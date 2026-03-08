"""
=============================================================
  PUNTO DE VENTA  —  Sistema local con SQLite
  Ejecutar:  python punto_de_venta.py
=============================================================
"""

import tkinter as tk
from tkinter import ttk, messagebox, font as tkfont
import sqlite3, os, datetime, hashlib, csv, unicodedata

# ── python-docx (inventario físico) — opcional ────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

# ──────────────────────────────────────────────────────────
#  BASE DE DATOS
# ──────────────────────────────────────────────────────────
DB_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ventas.db")

def get_conn():
    return sqlite3.connect(DB_FILE)

def _hash(texto):
    """Devuelve el SHA-256 hexadecimal de un texto. Único punto de hashing en todo el sistema."""
    return hashlib.sha256(texto.encode()).hexdigest()

def get_admin_hash():
    """Lee el hash de contraseña guardado en BD. Retorna None si aún no se ha creado."""
    with get_conn() as conn:
        row = conn.execute(
            "SELECT valor FROM configuracion WHERE clave = 'admin_hash'"
        ).fetchone()
    return row[0] if row else None

def set_admin_hash(nuevo_hash):
    """Guarda o actualiza el hash en BD (INSERT OR REPLACE)."""
    with get_conn() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO configuracion (clave, valor) VALUES ('admin_hash', ?)",
            (nuevo_hash,)
        )

def init_db():
    with get_conn() as conn:
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS configuracion (
                clave  TEXT PRIMARY KEY,
                valor  TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS productos (
                id        INTEGER PRIMARY KEY AUTOINCREMENT,
                codigo    TEXT    UNIQUE NOT NULL,
                nombre    TEXT    NOT NULL,
                precio    REAL    NOT NULL DEFAULT 0,
                costo     REAL    NOT NULL DEFAULT 0,
                stock     REAL    NOT NULL DEFAULT 0,
                categoria TEXT    DEFAULT 'General',
                a_granel  INTEGER NOT NULL DEFAULT 0,
                caducidad TEXT    DEFAULT NULL
            );

            CREATE TABLE IF NOT EXISTS categorias_productos (
                id      INTEGER PRIMARY KEY AUTOINCREMENT,
                nombre  TEXT UNIQUE COLLATE NOCASE NOT NULL
            );

            CREATE TABLE IF NOT EXISTS ventas (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                fecha      TEXT    NOT NULL,
                total      REAL    NOT NULL DEFAULT 0,
                cliente_id INTEGER DEFAULT NULL,
                FOREIGN KEY (cliente_id) REFERENCES clientes(id)
            );

            CREATE TABLE IF NOT EXISTS clientes (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                nombre        TEXT    NOT NULL,
                telefono      TEXT    DEFAULT '',
                email         TEXT    DEFAULT '',
                notas         TEXT    DEFAULT '',
                ultima_visita TEXT    DEFAULT NULL,
                total_compras REAL    NOT NULL DEFAULT 0,
                fecha_alta    TEXT    NOT NULL DEFAULT (datetime('now','localtime'))
            );

            CREATE TABLE IF NOT EXISTS categorias_contables (
                id         INTEGER PRIMARY KEY AUTOINCREMENT,
                nombre     TEXT UNIQUE NOT NULL,
                tipo       TEXT NOT NULL CHECK (tipo IN ('INGRESO','EGRESO')),
                naturaleza TEXT NOT NULL CHECK (naturaleza IN ('FIJO','VARIABLE','NO_APLICA')),
                activa     INTEGER NOT NULL DEFAULT 1
            );

            CREATE TABLE IF NOT EXISTS movimientos_caja (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                fecha       TEXT    NOT NULL DEFAULT (datetime('now','localtime')),
                tipo        TEXT    NOT NULL CHECK (tipo IN ('INGRESO','EGRESO')),
                subtipo     TEXT    NOT NULL CHECK (subtipo IN ('VENTA','GASTO','RETIRO_DUENO','DEPOSITO','AJUSTE','INGRESO_EXTRA')),
                categoria   TEXT    NOT NULL,
                concepto    TEXT    NOT NULL,
                metodo_pago TEXT    NOT NULL DEFAULT 'EFECTIVO',
                monto       REAL    NOT NULL CHECK (monto >= 0),
                referencia  TEXT    DEFAULT '',
                venta_id    INTEGER DEFAULT NULL,
                usuario     TEXT    DEFAULT 'admin',
                notas       TEXT    DEFAULT '',
                FOREIGN KEY (venta_id) REFERENCES ventas(id)
            );

            CREATE TABLE IF NOT EXISTS cortes_caja (
                id               INTEGER PRIMARY KEY AUTOINCREMENT,
                fecha_apertura   TEXT NOT NULL,
                fecha_cierre     TEXT DEFAULT NULL,
                saldo_inicial    REAL NOT NULL DEFAULT 0,
                ventas_efectivo  REAL NOT NULL DEFAULT 0,
                ingresos_extra   REAL NOT NULL DEFAULT 0,
                egresos          REAL NOT NULL DEFAULT 0,
                saldo_teorico    REAL NOT NULL DEFAULT 0,
                saldo_real       REAL DEFAULT NULL,
                diferencia       REAL DEFAULT NULL,
                estado           TEXT NOT NULL DEFAULT 'ABIERTO' CHECK (estado IN ('ABIERTO','CERRADO')),
                usuario_apertura TEXT DEFAULT 'admin',
                usuario_cierre   TEXT DEFAULT ''
            );

            CREATE TABLE IF NOT EXISTS detalle_venta (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                venta_id    INTEGER NOT NULL,
                producto_id INTEGER NOT NULL,
                nombre      TEXT    NOT NULL,
                precio      REAL    NOT NULL,
                costo       REAL    NOT NULL DEFAULT 0,
                cantidad    REAL    NOT NULL,
                subtotal    REAL    NOT NULL,
                ganancia    REAL    NOT NULL DEFAULT 0,
                es_granel   INTEGER NOT NULL DEFAULT 0,
                FOREIGN KEY (venta_id)    REFERENCES ventas(id),
                FOREIGN KEY (producto_id) REFERENCES productos(id)
            );
        """)
        # Migración: agregar columnas a BD existente sin perder datos
        for sql in [
            "ALTER TABLE productos ADD COLUMN costo REAL NOT NULL DEFAULT 0",
            "ALTER TABLE productos ADD COLUMN a_granel INTEGER NOT NULL DEFAULT 0",
            "ALTER TABLE productos ADD COLUMN caducidad TEXT DEFAULT NULL",
            "ALTER TABLE ventas ADD COLUMN cliente_id INTEGER DEFAULT NULL",
            "ALTER TABLE detalle_venta ADD COLUMN costo REAL NOT NULL DEFAULT 0",
            "ALTER TABLE detalle_venta ADD COLUMN ganancia REAL NOT NULL DEFAULT 0",
            "ALTER TABLE detalle_venta ADD COLUMN es_granel INTEGER NOT NULL DEFAULT 0",
            "ALTER TABLE clientes ADD COLUMN telefono TEXT DEFAULT ''",
            "ALTER TABLE clientes ADD COLUMN email TEXT DEFAULT ''",
            "ALTER TABLE clientes ADD COLUMN notas TEXT DEFAULT ''",
            "ALTER TABLE clientes ADD COLUMN ultima_visita TEXT DEFAULT NULL",
            "ALTER TABLE clientes ADD COLUMN total_compras REAL NOT NULL DEFAULT 0",
            "ALTER TABLE clientes ADD COLUMN fecha_alta TEXT NOT NULL DEFAULT (datetime('now','localtime'))",
            "ALTER TABLE movimientos_caja ADD COLUMN usuario TEXT DEFAULT 'admin'",
            "ALTER TABLE movimientos_caja ADD COLUMN notas TEXT DEFAULT ''",
        ]:
            try:
                conn.execute(sql)
            except Exception:
                pass  # La columna ya existe — ignorar

        cur_cat = conn.execute("SELECT COUNT(*) FROM categorias_contables")
        if cur_cat.fetchone()[0] == 0:
            conn.executemany(
                "INSERT INTO categorias_contables (nombre,tipo,naturaleza,activa) VALUES (?,?,?,1)",
                [
                    ("Ventas", "INGRESO", "NO_APLICA"),
                    ("Ingreso extra", "INGRESO", "NO_APLICA"),
                    ("Depósito", "INGRESO", "NO_APLICA"),
                    ("Ajuste de caja (+)", "INGRESO", "NO_APLICA"),
                    ("Compra inventario", "EGRESO", "VARIABLE"),
                    ("Renta", "EGRESO", "FIJO"),
                    ("Sueldos", "EGRESO", "FIJO"),
                    ("Servicios", "EGRESO", "FIJO"),
                    ("Transporte", "EGRESO", "VARIABLE"),
                    ("Mantenimiento", "EGRESO", "VARIABLE"),
                    ("Retiro dueño", "EGRESO", "NO_APLICA"),
                    ("Ajuste de caja (-)", "EGRESO", "NO_APLICA"),
                ]
            )
        cur = conn.execute("SELECT COUNT(*) FROM productos")
        if cur.fetchone()[0] == 0:
            conn.executemany(
                "INSERT INTO productos (codigo,nombre,precio,costo,stock,categoria) VALUES (?,?,?,?,?,?)",
                [
                    ("P001", "Refresco 600ml",  18.0, 12.0, 50, "Bebidas"),
                    ("P002", "Agua 500ml",       10.0,  6.0, 80, "Bebidas"),
                    ("P003", "Papas fritas",     15.0,  9.0, 30, "Botanas"),
                    ("P004", "Galletas",         12.0,  7.0, 40, "Botanas"),
                    ("P005", "Café americano",   25.0, 14.0, 20, "Cafetería"),
                ]
            )
        conn.execute(
            "INSERT OR IGNORE INTO categorias_productos (nombre) VALUES ('General')"
        )
        conn.execute(
            "INSERT OR IGNORE INTO categorias_productos (nombre)"
            " SELECT DISTINCT TRIM(categoria) FROM productos"
            " WHERE categoria IS NOT NULL AND TRIM(categoria) <> ''"
        )

def _generar_codigo_unico():
    with get_conn() as conn:
        row = conn.execute(
            "SELECT MAX(CAST(SUBSTR(codigo,2) AS INTEGER)) FROM productos"
            " WHERE codigo GLOB 'P[0-9][0-9][0-9][0-9]'"
        ).fetchone()
        maxn = row[0] or 0
        for i in range(1, 10001):
            cand = maxn + i
            if cand > 9999:
                break
            codigo = f"P{cand:04d}"
            existe = conn.execute(
                "SELECT 1 FROM productos WHERE codigo = ?", (codigo,)
            ).fetchone()
            if not existe:
                return codigo
    return f"P{datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')}"

# ──────────────────────────────────────────────────────────
#  GENERADOR DE INVENTARIO FÍSICO (.docx)
# ──────────────────────────────────────────────────────────

def _rgb(hex_str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    ex = tcPr.find(qn("w:shd"))
    if ex is not None: tcPr.remove(ex)
    tcPr.append(shd)

def _cell_borders(cell, color="C8E6C9", sz=4):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    ex = tcPr.find(qn("w:tcBorders"))
    if ex is not None: tcPr.remove(ex)
    bd = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        n = OxmlElement(f"w:{side}")
        n.set(qn("w:val"),"single"); n.set(qn("w:sz"),str(sz*8))
        n.set(qn("w:space"),"0"); n.set(qn("w:color"),color.lstrip("#").upper())
        bd.append(n)
    tcPr.append(bd)

def _cell_write(cell, text, bold=False, size=9, color="1B2E1C",
                align=None, bg=None):
    if align is None:
        align = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = align
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side,v in [("top","40"),("bottom","40"),("left","80"),("right","80")]:
        m = OxmlElement(f"w:{side}"); m.set(qn("w:w"),v); m.set(qn("w:type"),"dxa")
        mar.append(m)
    ex = tcPr.find(qn("w:tcMar"))
    if ex is not None: tcPr.remove(ex)
    tcPr.append(mar)
    r = p.add_run(str(text))
    r.bold = bold; r.font.name = "Arial"
    r.font.size = Pt(size); r.font.color.rgb = _rgb(color)
    if bg: _cell_bg(cell, bg)

def _col_width(table, col_idx, cm):
    for row in table.rows:
        tc = row.cells[col_idx]._tc; tcPr = tc.get_or_add_tcPr()
        w = OxmlElement("w:tcW")
        w.set(qn("w:w"), str(int(cm*567))); w.set(qn("w:type"),"dxa")
        ex = tcPr.find(qn("w:tcW"))
        if ex is not None: tcPr.remove(ex)
        tcPr.append(w)

def _tbl_fixed(table):
    tbl = table._tbl; tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); tbl.insert(0,tblPr)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"),"fixed")
    ex = tblPr.find(qn("w:tblLayout"))
    if ex is not None: tblPr.remove(ex)
    tblPr.append(lay)

def generar_docx_inventario(productos, ruta_salida):
    """
    Genera un .docx de levantamiento de inventario físico.
    productos: list[tuple] → (id, codigo, nombre, precio, costo, stock, categoria)
    ruta_salida: str → ruta completa donde guardar el archivo
    Propaga excepciones al llamador para que las maneje con messagebox.
    """
    CLR = {
        "hdr":     "2E7D32", "sub":  "A5D6A7", "fila":  "F1F8E9",
        "bosque":  "39542C", "hier": "48872B", "borde": "C8E6C9",
        "borde_o": "39542C", "txt":  "1B2E1C", "mut":   "555555",
        "rojo":    "C62828", "bco":  "FFFFFF",
    }
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = 1          # landscape
    sec.page_width  = Cm(29.7); sec.page_height = Cm(21.0)
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(1.5)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Encabezado ─────────────────────────────────────────
    t = doc.add_paragraph()
    r = t.add_run("HOJA DE LEVANTAMIENTO DE INVENTARIO FÍSICO")
    r.bold = True; r.font.size = Pt(16); r.font.name = "Arial"
    r.font.color.rgb = _rgb(CLR["bosque"])
    s = doc.add_paragraph()
    rs = s.add_run("Sistema: Punto de Venta  •  Base de datos: ventas.db  •  Tabla: productos")
    rs.font.size = Pt(9); rs.font.name = "Arial"; rs.font.color.rgb = _rgb(CLR["hier"])
    s.paragraph_format.space_after = Pt(4)
    hr = doc.add_paragraph(); hr.paragraph_format.space_after = Pt(6)
    pPr = hr._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom"); bot.set(qn("w:val"),"single")
    bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"1")
    bot.set(qn("w:color"),CLR["bosque"]); pBdr.append(bot); pPr.append(pBdr)

    # ── Datos del levantamiento ────────────────────────────
    fecha_hoy = datetime.datetime.now().strftime("%d / %m / %Y")
    it = doc.add_table(rows=2, cols=4); _tbl_fixed(it)
    it.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx,(lbl,val) in enumerate([
        ("Responsable del conteo:",""),("Fecha del conteo:",fecha_hoy),
        ("Supervisor / Autoriza:",""),("Folio / Consecutivo:",""),
    ]):
        ri,ci = idx//2, (idx%2)*2
        _cell_write(it.cell(ri,ci),lbl,bold=True,size=8,color=CLR["bosque"],
                    align=WD_ALIGN_PARAGRAPH.LEFT,bg=CLR["fila"])
        _cell_write(it.cell(ri,ci+1),val,size=9,align=WD_ALIGN_PARAGRAPH.LEFT,bg=CLR["bco"])
        for c in (it.cell(ri,ci), it.cell(ri,ci+1)): _cell_borders(c,CLR["borde"])
    for i,w in enumerate([4.5,5.5,4.5,5.5]): _col_width(it,i,w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Tabla de inventario ────────────────────────────────
    COLS = [("#",1.0),("Código",2.0),("Categoría",3.0),("Nombre del Producto",6.0),
            ("Costo ($)",2.2),("Precio ($)",2.2),("Stock\nSistema",2.2),
            ("Stock\nFísico",2.2),("Diferencia",2.2),("Observaciones",3.7)]
    n = len(COLS)
    tp = doc.add_paragraph(); tp.paragraph_format.space_before = Pt(2)
    tp.paragraph_format.space_after = Pt(2)
    rt2 = tp.add_run("INVENTARIO DE PRODUCTOS")
    rt2.bold=True; rt2.font.size=Pt(10); rt2.font.name="Arial"
    rt2.font.color.rgb=_rgb(CLR["bosque"])
    total_stock = sum(p[5] for p in productos)
    rp = doc.add_paragraph(); rp.paragraph_format.space_after = Pt(3)
    rr2 = rp.add_run(
        f"Total de registros: {len(productos)} productos  •  "
        f"Stock total en sistema: {total_stock} unidades")
    rr2.font.size=Pt(8.5); rr2.font.name="Arial"; rr2.font.color.rgb=_rgb(CLR["hier"])

    inv = doc.add_table(rows=1, cols=n); _tbl_fixed(inv)
    inv.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,(lbl,_) in enumerate(COLS):
        _cell_write(inv.rows[0].cells[i],lbl,bold=True,size=8.5,
                    color=CLR["bco"],bg=CLR["hdr"])
        _cell_borders(inv.rows[0].cells[i],CLR["borde_o"],5)

    categorias = list(dict.fromkeys(p[6] for p in productos))
    for cat in categorias:
        sep = inv.add_row(); mc = sep.cells[0]
        for j in range(1,n): mc = mc.merge(sep.cells[j])
        _cell_write(mc,f"  ▶  {cat.upper()}",bold=True,size=9,color=CLR["bosque"],
                    align=WD_ALIGN_PARAGRAPH.LEFT,bg=CLR["sub"])
        _cell_borders(mc,CLR["borde_o"])
        for fi,prod in enumerate([p for p in productos if p[6]==cat]):
            pid,codigo,nombre,precio,costo,stock,categoria = prod
            bg = CLR["bco"] if fi%2==0 else CLR["fila"]
            sc = CLR["rojo"] if stock<=5 else CLR["txt"]
            dr = inv.add_row()
            vals = [(str(pid),CLR["txt"],False),(codigo,CLR["txt"],False),
                    (categoria,CLR["mut"],False),(nombre,CLR["txt"],False),
                    (f"${costo:.2f}",CLR["mut"],False),(f"${precio:.2f}",CLR["txt"],False),
                    (str(stock),sc,stock<=5),("","",False),("","",False),("","",False)]
            alns = [WD_ALIGN_PARAGRAPH.CENTER]*3+[WD_ALIGN_PARAGRAPH.LEFT]+\
                   [WD_ALIGN_PARAGRAPH.CENTER]*5+[WD_ALIGN_PARAGRAPH.LEFT]
            for j,(v,col,bld) in enumerate(vals):
                _cell_write(dr.cells[j],v,bold=bld,size=9,
                            color=col if col else CLR["txt"],align=alns[j],bg=bg)
                _cell_borders(dr.cells[j],CLR["borde"])

    tr = inv.add_row(); tl = tr.cells[0]
    for j in range(1,6): tl = tl.merge(tr.cells[j])
    _cell_write(tl,"TOTAL DE UNIDADES EN SISTEMA:",bold=True,size=9.5,
                color=CLR["bco"],bg=CLR["bosque"],align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_borders(tl,CLR["borde_o"],5)
    _cell_write(tr.cells[6],str(total_stock),bold=True,size=10,
                color=CLR["bco"],bg=CLR["bosque"]); _cell_borders(tr.cells[6],CLR["borde_o"],5)
    for j in range(7,n):
        _cell_write(tr.cells[j],"",bg=CLR["fila"]); _cell_borders(tr.cells[j],CLR["borde"])
    for i,(_,w) in enumerate(COLS): _col_width(inv,i,w)

    # ── Instrucciones ──────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    ti = doc.add_paragraph(); ti.paragraph_format.space_after = Pt(3)
    ri2 = ti.add_run("INSTRUCCIONES PARA EL LEVANTAMIENTO")
    ri2.bold=True; ri2.font.size=Pt(10); ri2.font.name="Arial"
    ri2.font.color.rgb=_rgb(CLR["bosque"])
    inst = [("1.","Cuenta físicamente cada producto y escribe la cantidad real en STOCK FÍSICO."),
            ("2.","DIFERENCIA = Stock Sistema − Stock Físico. Negativo = faltante; positivo = sobrante."),
            ("3.","Registra en OBSERVACIONES productos dañados, vencidos o con discrepancia sin explicación."),
            ("4.","Al terminar, responsable y supervisor deben firmar al pie de este documento."),
            ("5.","Conserva este impreso y actualiza los valores en el sistema si corresponde.")]
    itbl = doc.add_table(rows=len(inst),cols=2); _tbl_fixed(itbl)
    for i,(num,txt) in enumerate(inst):
        bg = CLR["fila"] if i%2==0 else CLR["bco"]
        _cell_write(itbl.rows[i].cells[0],num,bold=True,size=9,color=CLR["bosque"],bg=bg)
        _cell_write(itbl.rows[i].cells[1],txt,size=9,color=CLR["txt"],
                    align=WD_ALIGN_PARAGRAPH.LEFT,bg=bg)
        for c in itbl.rows[i].cells: _cell_borders(c,CLR["borde"])
    _col_width(itbl,0,0.8); _col_width(itbl,1,25.9)

    # ── Firmas ─────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    ftbl = doc.add_table(rows=2,cols=3); _tbl_fixed(ftbl)
    ftbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_border = {"top":"none","left":"none","bottom":"none","right":"none"}
    for j in range(3):
        c = ftbl.rows[0].cells[j]; _cell_write(c,"",bg=CLR["bco"])
        tc=c._tc; tcPr=tc.get_or_add_tcPr(); bd=OxmlElement("w:tcBorders")
        for side,val in _no_border.items():
            nd=OxmlElement(f"w:{side}"); nd.set(qn("w:val"),val); bd.append(nd)
        tcPr.append(bd)
    for j,etiq in enumerate(["Responsable del conteo","","Supervisor / Autoriza"]):
        c=ftbl.rows[1].cells[j]; tc=c._tc; tcPr=tc.get_or_add_tcPr()
        bd=OxmlElement("w:tcBorders")
        for side in ("left","right","bottom"):
            nd=OxmlElement(f"w:{side}"); nd.set(qn("w:val"),"none"); bd.append(nd)
        top=OxmlElement("w:top")
        if etiq:
            top.set(qn("w:val"),"single"); top.set(qn("w:sz"),"8")
            top.set(qn("w:color"),CLR["bosque"])
        else:
            top.set(qn("w:val"),"none")
        bd.append(top); tcPr.append(bd)
        if etiq:
            p2=c.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rr3=p2.add_run(etiq); rr3.bold=True; rr3.font.size=Pt(9)
            rr3.font.name="Arial"; rr3.font.color.rgb=_rgb(CLR["bosque"])
    _col_width(ftbl,0,11.5); _col_width(ftbl,1,3.5); _col_width(ftbl,2,11.5)

    doc.save(ruta_salida)

# ──────────────────────────────────────────────────────────
#  COLORES Y ESTILO
# ──────────────────────────────────────────────────────────
C = {
    "bg":        "#4CBB17",
    "panel":     "#ffffff",
    "card":      "#ffffff",
    "border":    "#e6f3e8",
    "accent":    "#48872B",
    "accent2":   "#39542C",
    "green":     "#293325",
    "red":       "#e74c3c",
    "yellow":    "#f39c12",
    "text":      "#062b00",
    "muted":     "#6b786e",
    "white":     "#ffffff",
    "hover":     "#a8e39a",
}

# ──────────────────────────────────────────────────────────
#  APLICACIÓN PRINCIPAL
# ──────────────────────────────────────────────────────────
class PuntoDeVenta(tk.Tk):
    def __init__(self):
        super().__init__()
        init_db()
        self.title("Punto de Venta")
        self.geometry("1200x750")
        self.minsize(900, 600)
        self.configure(bg=C["bg"])
        self.option_add("*Font", "Courier 12")
        self.carrito = []
        self._alertas_caducidad_notificadas = set()
        self._clientes_cache = []
        self._cliente_id_por_opcion = {"Público general": None}
        self._categorias_contables_cache = []
        self._categorias_producto_cache = []
        self._build_ui()
        self._cargar_productos()
        self._cargar_clientes_en_venta()
        self.after(200, self._verificar_contrasena_inicial)

    # ── UI principal ──────────────────────────────────────
    def _build_ui(self):
        header = tk.Frame(self, bg=C["bg"], pady=0)
        header.pack(fill="x", padx=20, pady=(14, 0))

        tk.Label(header, text="●", fg=C["accent"], bg=C["bg"],
                 font=("Courier", 20)).pack(side="left")
        tk.Label(header, text="  PUNTO DE VENTA", fg=C["text"], bg=C["bg"],
                 font=("Courier", 16, "bold")).pack(side="left")

        self.nav_btns = {}
        nav_frame = tk.Frame(header, bg=C["bg"])
        nav_frame.pack(side="right")
        for label, cmd in [("🛒  Ventas", self._show_ventas),
                           ("📦  Productos", self._show_productos),
                           ("👥  CRM", self._show_crm),
                           ("💰  Contabilidad", self._show_contabilidad),
                           ("📊  Historial", self._show_historial)]:
            b = tk.Button(nav_frame, text=label, bg=C["panel"], fg=C["muted"],
                          bd=0, padx=16, pady=6, cursor="hand2",
                          font=("Courier", 10), activebackground=C["hover"],
                          activeforeground=C["text"], command=cmd)
            b.pack(side="left", padx=3)
            self.nav_btns[label] = b

        sep = tk.Frame(self, bg=C["border"], height=1)
        sep.pack(fill="x", padx=20, pady=10)

        self.content = tk.Frame(self, bg=C["bg"])
        self.content.pack(fill="both", expand=True, padx=20, pady=(0,16))

        self.pages = {}
        self._build_page_ventas()
        self._build_page_productos()
        self._build_page_crm()
        self._build_page_contabilidad()
        self._build_page_historial()
        self._show_ventas()

    def _show_page(self, name):
        for p in self.pages.values():
            p.pack_forget()
        self.pages[name].pack(fill="both", expand=True)
        labels = {"ventas": "🛒  Ventas", "productos": "📦  Productos",
                  "crm": "👥  CRM", "contabilidad": "💰  Contabilidad", "historial": "📊  Historial"}
        for k, b in self.nav_btns.items():
            b.config(fg=C["accent"] if k == labels[name] else C["muted"],
                     bg=C["card"] if k == labels[name] else C["panel"])

    def _show_ventas(self):    self._show_page("ventas")
    def _show_productos(self):
        self._refrescar_categorias_producto()
        self._cargar_tabla_productos()
        self._show_page("productos")
        if not getattr(self, "_editing_id", None):
            cod = self._prod_entries["e_codigo"].get().strip()
            if not cod:
                self._autocodigo()
    def _show_crm(self):
        self._cargar_tabla_clientes()
        self._show_page("crm")
    def _show_contabilidad(self):
        self._cargar_contabilidad_vista()
        self._show_page("contabilidad")
    def _show_historial(self): self._cargar_historial(); self._show_page("historial")

    # ══════════════════════════════════════════════════════
    #  PÁGINA: VENTAS
    # ══════════════════════════════════════════════════════
    def _build_page_ventas(self):
        page = tk.Frame(self.content, bg=C["bg"])
        self.pages["ventas"] = page

        left = tk.Frame(page, bg=C["bg"])
        left.pack(side="left", fill="both", expand=True, padx=(0,12))

        search_card = tk.Frame(left, bg=C["card"], padx=12, pady=10)
        search_card.pack(fill="x", pady=(0,10))
        tk.Label(search_card, text="BUSCAR PRODUCTO", fg=C["muted"],
                 bg=C["card"], font=("Courier", 9)).pack(anchor="w")

        sf = tk.Frame(search_card, bg=C["border"], pady=1)
        sf.pack(fill="x", pady=(4,0))
        inner = tk.Frame(sf, bg=C["panel"])
        inner.pack(fill="x")

        tk.Label(inner, text="⌕", fg=C["accent"], bg=C["panel"],
                 font=("Courier", 14), padx=8).pack(side="left")

        self.sv_busqueda = tk.StringVar()
        self.sv_busqueda.trace_add("write", lambda *a: self._filtrar_productos())
        entry = tk.Entry(inner, textvariable=self.sv_busqueda,
                         bg=C["panel"], fg=C["text"], insertbackground=C["text"],
                         bd=0, font=("Courier", 12), highlightthickness=0)
        entry.pack(side="left", fill="x", expand=True, ipady=8)
        entry.focus()
        entry.bind("<Return>", lambda e: self._agregar_primero_al_carrito())
        entry.bind("<Down>", lambda e: self._focus_tabla())

        cols = ("codigo", "nombre", "precio", "stock")
        heads = ("Código", "Nombre", "Precio", "Stock")
        frame_t = tk.Frame(left, bg=C["bg"])
        frame_t.pack(fill="both", expand=True)

        style = ttk.Style()
        style.theme_use("clam")
        style.configure("POS.Treeview",
            background=C["card"], fieldbackground=C["card"],
            foreground=C["text"], rowheight=38,
            font=("Courier", 12), borderwidth=0)
        style.configure("POS.Treeview.Heading",
            background=C["border"], foreground=C["muted"],
            font=("Courier", 11, "bold"), relief="flat")
        style.map("POS.Treeview",
            background=[("selected", C["accent2"])],
            foreground=[("selected", C["white"])])

        self.tabla_busq = ttk.Treeview(frame_t, columns=cols, show="headings",
                                       style="POS.Treeview", selectmode="browse")
        widths = [80, 260, 90, 70]
        for c, h, w in zip(cols, heads, widths):
            self.tabla_busq.heading(c, text=h, anchor="center")
            self.tabla_busq.column(c, width=w, anchor="center", stretch=True)

        sb = ttk.Scrollbar(frame_t, orient="vertical",
               command=self.tabla_busq.yview)
        self.tabla_busq.configure(yscrollcommand=sb.set)
        self.tabla_busq.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.tabla_busq.bind("<Double-1>", lambda e: self._agregar_seleccionado())
        self.tabla_busq.bind("<Return>", lambda e: self._agregar_seleccionado())

        btn_add = tk.Button(left, text="＋  Agregar al carrito  (↵ Enter)",
                            bg=C["accent"], fg=C["white"], bd=0,
                            font=("Courier", 11, "bold"), pady=10, cursor="hand2",
                            activebackground="#3a7de0",
                            command=self._agregar_seleccionado)
        btn_add.pack(fill="x", pady=(8,0))

        right = tk.Frame(page, bg=C["panel"], width=360, padx=14, pady=14)
        right.pack(side="right", fill="y")
        right.pack_propagate(False)

        tk.Label(right, text="CARRITO DE VENTA", fg=C["muted"],
                 bg=C["panel"], font=("Courier", 9, "bold")).pack(anchor="w")

        cliente_f = tk.Frame(right, bg=C["panel"])
        cliente_f.pack(fill="x", pady=(8, 4))
        tk.Label(cliente_f, text="Cliente (opcional)", fg=C["muted"], bg=C["panel"],
                 font=("Courier", 9)).pack(anchor="w")
        tk.Label(cliente_f, text="Buscar cliente", fg=C["muted"], bg=C["panel"],
                 font=("Courier", 9)).pack(anchor="w", pady=(4, 0))
        self.sv_cliente_busqueda_venta = tk.StringVar(value="")
        self.sv_cliente_busqueda_venta.trace_add(
            "write", lambda *a: self._filtrar_clientes_en_venta()
        )
        tk.Entry(
            cliente_f,
            textvariable=self.sv_cliente_busqueda_venta,
            bg=C["card"], fg=C["text"], insertbackground=C["text"],
            bd=0, font=("Courier", 10), highlightthickness=1,
            highlightbackground=C["border"]
        ).pack(fill="x", ipady=4, pady=(2, 4))
        self.frame_lista_clientes_venta = tk.Frame(cliente_f, bg=C["panel"])
        self.lst_clientes_venta = tk.Listbox(
            self.frame_lista_clientes_venta,
            bg=C["card"], fg=C["text"], selectbackground=C["accent2"],
            bd=0, font=("Courier", 9), activestyle="none",
            highlightthickness=1, highlightbackground=C["border"], height=4
        )
        self.lst_clientes_venta.pack(fill="x")
        self.lst_clientes_venta.bind("<<ListboxSelect>>", self._seleccionar_cliente_lista_venta)
        self.lst_clientes_venta.bind("<Return>", self._seleccionar_cliente_lista_venta)
        self.frame_lista_clientes_venta.pack_forget()
        self.sv_cliente_venta = tk.StringVar(value="Público general")
        self.cmb_cliente_venta = ttk.Combobox(
            cliente_f,
            textvariable=self.sv_cliente_venta,
            state="readonly",
            font=("Courier", 10)
        )
        self.cmb_cliente_venta.bind(
            "<<ComboboxSelected>>", lambda _e: self.sv_cliente_busqueda_venta.set("")
        )
        self.cmb_cliente_venta.pack(fill="x", ipady=4, pady=(2, 0))
        self.cmb_cliente_venta["values"] = ("Público general",)
        self.cmb_cliente_venta.current(0)

        tk.Frame(right, bg=C["border"], height=1).pack(fill="x", pady=8)

        cart_frame = tk.Frame(right, bg=C["panel"])
        cart_frame.pack(fill="both", expand=True)

        self.lista_carrito = tk.Listbox(cart_frame, bg=C["card"], fg=C["text"],
                                        selectbackground=C["accent2"], bd=0,
                                        font=("Courier", 10), activestyle="none",
                                        highlightthickness=0, relief="flat")
        sb2 = ttk.Scrollbar(cart_frame, orient="vertical",
                             command=self.lista_carrito.yview)
        self.lista_carrito.configure(yscrollcommand=sb2.set)
        self.lista_carrito.pack(side="left", fill="both", expand=True)
        sb2.pack(side="right", fill="y")

        tk.Frame(right, bg=C["border"], height=1).pack(fill="x", pady=8)

        total_frame = tk.Frame(right, bg=C["panel"])
        total_frame.pack(fill="x")
        tk.Label(total_frame, text="TOTAL", fg=C["muted"],
                 bg=C["panel"], font=("Courier", 10)).pack(side="left")
        self.lbl_total = tk.Label(total_frame, text="$0.00", fg=C["green"],
                                  bg=C["panel"], font=("Courier", 22, "bold"))
        self.lbl_total.pack(side="right")

        btn_quitar = tk.Button(right, text="✕  Quitar seleccionado",
                               bg=C["card"], fg=C["red"], bd=0,
                               font=("Courier", 10), pady=8, cursor="hand2",
                               activebackground=C["hover"],
                               command=self._quitar_del_carrito)
        btn_quitar.pack(fill="x", pady=(10,4))

        btn_limpiar = tk.Button(right, text="⟳  Limpiar carrito",
                                bg=C["card"], fg=C["yellow"], bd=0,
                                font=("Courier", 10), pady=8, cursor="hand2",
                                activebackground=C["hover"],
                                command=self._limpiar_carrito)
        btn_limpiar.pack(fill="x", pady=4)

        btn_cobrar = tk.Button(right, text="✔  COBRAR VENTA",
                               bg=C["green"], fg=C["white"], bd=0,
                               font=("Courier", 13, "bold"), pady=14, cursor="hand2",
                               activebackground="#27ae60",
                               command=self._cobrar_venta)
        btn_cobrar.pack(fill="x", pady=(8,0))

    def _fmt_unidades(self, valor, es_granel=False, con_unidad=False):
        try:
            num = float(valor)
        except (TypeError, ValueError):
            return "0"
        if es_granel:
            txt = f"{num:.3f}".rstrip("0").rstrip(".")
            if not txt:
                txt = "0"
            return f"{txt} kg" if con_unidad else txt
        return str(int(round(num)))

    def _texto_busqueda(self, texto):
        base = (texto or "").strip().lower()
        nfkd = unicodedata.normalize("NFD", base)
        return "".join(c for c in nfkd if unicodedata.category(c) != "Mn")

    def _estado_caducidad(self, fecha_txt):
        fecha_txt = (fecha_txt or "").strip()
        if not fecha_txt:
            return False, None
        try:
            fecha = datetime.date.fromisoformat(fecha_txt)
        except ValueError:
            return False, None
        dias = (fecha - datetime.date.today()).days
        return dias <= 30, dias

    def _texto_estado_caducidad(self, dias):
        if dias is None:
            return "Sin fecha"
        if dias < 0:
            return f"Vencido ({abs(dias)}d)"
        if dias <= 30:
            return f"CRITICO ({dias}d)"
        return f"OK ({dias}d)"

    def _clave_prioridad_salida(self, fecha_txt, nombre):
        es_critico, dias = self._estado_caducidad(fecha_txt)
        if es_critico:
            return (0, dias, nombre.lower())
        if dias is not None:
            return (1, dias, nombre.lower())
        return (2, 999999, nombre.lower())

    def _avisar_caducidades_criticas(self, productos):
        criticos_nuevos = []
        for prod in productos:
            pid, codigo, nombre, _precio, _costo, _stock, _a_granel, caducidad, *_extra = prod
            es_critico, dias = self._estado_caducidad(caducidad)
            if not es_critico:
                continue
            clave_alerta = (pid, caducidad)
            if clave_alerta in self._alertas_caducidad_notificadas:
                continue
            self._alertas_caducidad_notificadas.add(clave_alerta)
            criticos_nuevos.append((codigo, nombre, dias, caducidad))

        if not criticos_nuevos:
            return

        lineas = []
        for codigo, nombre, dias, cad in criticos_nuevos[:12]:
            estado = "VENCIDO" if dias < 0 else f"caduca en {dias} día(s)"
            lineas.append(f"- {codigo} | {nombre} | {cad} | {estado}")
        extra = len(criticos_nuevos) - 12
        if extra > 0:
            lineas.append(f"... y {extra} producto(s) adicional(es).")

        messagebox.showwarning(
            "⚠ Productos críticos por caducidad",
            "Los siguientes productos están en estado CRÍTICO (menos de 1 mes para caducar) "
            "y deben salir con prioridad:\n\n" + "\n".join(lineas),
            parent=self
        )

    def _pedir_cantidad_granel(self, nombre, precio, stock_disponible):
        dlg = tk.Toplevel(self)
        dlg.title("Producto a granel")
        dlg.configure(bg=C["card"])
        dlg.resizable(False, False)
        dlg.transient(self)
        dlg.grab_set()
        dlg.focus_set()
        self.update_idletasks()
        x = self.winfo_x() + (self.winfo_width()  // 2) - 235
        y = self.winfo_y() + (self.winfo_height() // 2) - 170
        dlg.geometry(f"470x340+{x}+{y}")

        tk.Label(dlg, text="VENTA A GRANEL", fg=C["accent2"], bg=C["card"],
                 font=("Courier", 12, "bold")).pack(pady=(16, 4))
        tk.Label(dlg, text=nombre, fg=C["text"], bg=C["card"],
                 font=("Courier", 10, "bold"), wraplength=420,
                 justify="center").pack(pady=(0, 6))
        tk.Label(
            dlg,
            text=f"Precio por kg: ${precio:.2f}   •   Disponible: {self._fmt_unidades(stock_disponible, True, True)}",
            fg=C["muted"], bg=C["card"], font=("Courier", 9)
        ).pack()
        tk.Label(dlg,
                 text="Captura cantidad o importe. El otro campo se calcula automáticamente.",
                 fg=C["muted"], bg=C["card"], font=("Courier", 9)).pack(pady=(4, 10))

        body = tk.Frame(dlg, bg=C["card"])
        body.pack(fill="x", padx=20)

        sv_cantidad = tk.StringVar(value="")
        sv_importe = tk.StringVar(value="")
        syncing = {"on": False}
        resultado = {"cantidad": None}

        def _parse_float(texto):
            t = texto.strip().replace(",", ".")
            if not t:
                return None
            try:
                return float(t)
            except ValueError:
                return None

        def _set_importe(*_):
            if syncing["on"]:
                return
            cantidad = _parse_float(sv_cantidad.get())
            if cantidad is None:
                return
            syncing["on"] = True
            sv_importe.set(f"{cantidad * precio:.2f}")
            syncing["on"] = False

        def _set_cantidad(*_):
            if syncing["on"] or precio <= 0:
                return
            importe = _parse_float(sv_importe.get())
            if importe is None:
                return
            cantidad = round(importe / precio, 3)
            syncing["on"] = True
            sv_cantidad.set(self._fmt_unidades(cantidad, True))
            syncing["on"] = False

        tk.Label(body, text="Cantidad vendida (kg)", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).pack(anchor="w")
        e_cantidad = tk.Entry(body, textvariable=sv_cantidad,
                              bg=C["panel"], fg=C["text"],
                              insertbackground=C["text"], bd=0,
                              font=("Courier", 12), highlightthickness=1,
                              highlightbackground=C["border"])
        e_cantidad.pack(fill="x", ipady=7, pady=(2, 10))

        tk.Label(body, text="Importe a cobrar ($)", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).pack(anchor="w")
        e_importe = tk.Entry(body, textvariable=sv_importe,
                             bg=C["panel"], fg=C["text"],
                             insertbackground=C["text"], bd=0,
                             font=("Courier", 12), highlightthickness=1,
                             highlightbackground=C["border"])
        e_importe.pack(fill="x", ipady=7)
        e_cantidad.focus()

        sv_cantidad.trace_add("write", _set_importe)
        sv_importe.trace_add("write", _set_cantidad)

        lbl_error = tk.Label(dlg, text="", fg=C["red"], bg=C["card"],
                             font=("Courier", 9))
        lbl_error.pack(pady=(8, 4))

        def _confirmar():
            cantidad = _parse_float(sv_cantidad.get())
            if cantidad is None or cantidad <= 0:
                lbl_error.config(text="Ingresa una cantidad válida mayor a 0.")
                return
            if cantidad > stock_disponible + 1e-9:
                lbl_error.config(
                    text=f"Stock insuficiente. Máximo: {self._fmt_unidades(stock_disponible, True, True)}"
                )
                return
            resultado["cantidad"] = round(cantidad, 3)
            dlg.destroy()

        def _cancelar():
            dlg.destroy()

        btns = tk.Frame(dlg, bg=C["card"])
        btns.pack(fill="x", padx=20, pady=(2, 12))
        tk.Button(btns, text="✔ Confirmar", bg=C["green"], fg=C["white"], bd=0,
                  font=("Courier", 10, "bold"), pady=8, cursor="hand2",
                  activebackground="#27ae60",
                  command=_confirmar).pack(side="left", fill="x", expand=True, padx=(0, 4))
        tk.Button(btns, text="✕ Cancelar", bg=C["panel"], fg=C["muted"], bd=0,
                  font=("Courier", 10), pady=8, cursor="hand2",
                  command=_cancelar).pack(side="left", fill="x", expand=True)

        dlg.bind("<Return>", lambda e: _confirmar())
        dlg.bind("<Escape>", lambda e: _cancelar())
        dlg.protocol("WM_DELETE_WINDOW", _cancelar)
        self.wait_window(dlg)
        return resultado["cantidad"]

    # ── Lógica de búsqueda ────────────────────────────────
    def _cargar_productos(self):
        self._productos_cache = []
        with get_conn() as conn:
            rows = conn.execute(
                "SELECT id,codigo,nombre,precio,costo,stock,a_granel,caducidad,categoria"
                " FROM productos ORDER BY nombre"
            ).fetchall()
        self._productos_cache = sorted(
            rows,
            key=lambda p: self._clave_prioridad_salida(p[7], p[2])
        )
        self._avisar_caducidades_criticas(self._productos_cache)
        self._filtrar_productos()

    def _filtrar_productos(self):
        q = self._texto_busqueda(self.sv_busqueda.get())
        for row in self.tabla_busq.get_children():
            self.tabla_busq.delete(row)
        for prod in self._productos_cache:
            pid, codigo, nombre, precio, costo, stock, a_granel, caducidad, categoria = prod
            codigo_q = self._texto_busqueda(codigo)
            nombre_q = self._texto_busqueda(nombre)
            categoria_q = self._texto_busqueda(categoria)
            if q in codigo_q or q in nombre_q or q in categoria_q:
                es_critico, _dias = self._estado_caducidad(caducidad)
                tags = []
                if stock <= 5:
                    tags.append("low")
                if es_critico:
                    tags.append("critical")
                stock_txt = self._fmt_unidades(stock, bool(a_granel), bool(a_granel))
                nombre_txt = f"⚠ {nombre}" if es_critico else nombre
                self.tabla_busq.insert("", "end",
                    values=(codigo, nombre_txt, f"${precio:.2f}", stock_txt),
                    iid=str(pid), tags=tuple(tags))
        self.tabla_busq.tag_configure("low", foreground=C["yellow"])
        self.tabla_busq.tag_configure("critical", foreground=C["red"])

    def _focus_tabla(self):
        children = self.tabla_busq.get_children()
        if children:
            self.tabla_busq.selection_set(children[0])
            self.tabla_busq.focus(children[0])
            self.tabla_busq.focus_set()

    def _agregar_primero_al_carrito(self):
        children = self.tabla_busq.get_children()
        if children:
            self.tabla_busq.selection_set(children[0])
            self._agregar_seleccionado()

    def _agregar_seleccionado(self):
        sel = self.tabla_busq.selection()
        if not sel:
            self._agregar_primero_al_carrito()
            return
        pid = int(sel[0])
        prod = next((p for p in self._productos_cache if p[0] == pid), None)
        if not prod:
            return
        pid, codigo, nombre, precio, costo, stock, a_granel, caducidad, _categoria = prod
        es_granel = bool(a_granel)
        if stock <= 0:
            messagebox.showwarning("Sin stock",
                f'"{nombre}" no tiene stock disponible.', parent=self)
            return
        item_existente = next((i for i in self.carrito if i["id"] == pid), None)
        if es_granel:
            ya_en_carrito = item_existente["cantidad"] if item_existente else 0
            stock_disponible = round(stock - ya_en_carrito, 3)
            if stock_disponible <= 0:
                messagebox.showwarning("Stock insuficiente",
                    f'Stock máximo: {self._fmt_unidades(stock, True, True)}', parent=self)
                return
            cantidad = self._pedir_cantidad_granel(nombre, precio, stock_disponible)
            if cantidad is None:
                return
            if item_existente:
                item_existente["cantidad"] = round(item_existente["cantidad"] + cantidad, 3)
            else:
                self.carrito.append({
                    "id": pid, "codigo": codigo, "nombre": nombre,
                    "precio": precio, "costo": costo, "cantidad": cantidad,
                    "stock": stock, "es_granel": True, "caducidad": caducidad
                })
            self._refresh_carrito()
            self.sv_busqueda.set("")
            return

        for item in self.carrito:
            if item["id"] == pid:
                if item["cantidad"] >= stock:
                    messagebox.showwarning("Stock insuficiente",
                        f'Stock máximo: {self._fmt_unidades(stock)}', parent=self)
                    return
                item["cantidad"] += 1
                self._refresh_carrito()
                self.sv_busqueda.set("")
                return
        self.carrito.append({"id": pid, "codigo": codigo, "nombre": nombre,
                              "precio": precio, "costo": costo,
                              "cantidad": 1, "stock": stock, "es_granel": False,
                              "caducidad": caducidad})
        self._refresh_carrito()
        self.sv_busqueda.set("")

    # ── Carrito ───────────────────────────────────────────
    def _refresh_carrito(self):
        self.lista_carrito.delete(0, "end")
        total = 0
        for i, item in enumerate(self.carrito):
            sub = item["precio"] * item["cantidad"]
            total += sub
            es_granel = bool(item.get("es_granel", False))
            cant_txt = self._fmt_unidades(item["cantidad"], es_granel, es_granel)
            line = f"  {item['nombre'][:22]:<22}  x{cant_txt}  ${sub:.2f}"
            self.lista_carrito.insert("end", line)
            if i % 2 == 0:
                self.lista_carrito.itemconfig(i, bg=C["card"])
            else:
                self.lista_carrito.itemconfig(i, bg=C["hover"])
        self.lbl_total.config(text=f"${total:.2f}")

    def _quitar_del_carrito(self):
        sel = self.lista_carrito.curselection()
        if not sel:
            return
        idx = sel[0]
        self.carrito.pop(idx)
        self._refresh_carrito()

    def _limpiar_carrito(self):
        if not self.carrito:
            return
        if messagebox.askyesno("Limpiar", "¿Vaciar el carrito?", parent=self):
            self.carrito.clear()
            self._refresh_carrito()

    def _cobrar_venta(self):
        if not self.carrito:
            messagebox.showinfo("Carrito vacío", "Agrega productos antes de cobrar.",
                                parent=self)
            return
        total = sum(i["precio"] * i["cantidad"] for i in self.carrito)
        confirm = messagebox.askyesno("Confirmar venta",
            f"¿Registrar venta por ${total:.2f}?", parent=self)
        if not confirm:
            return
        cliente_id = self._cliente_id_venta_actual()
        venta_id = None
        fecha = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        conn = None
        try:
            conn = get_conn()
            conn.execute("BEGIN")
            cur = conn.execute(
                "INSERT INTO ventas (fecha,total,cliente_id) VALUES (?,?,?)",
                (fecha, total, cliente_id)
            )
            venta_id = cur.lastrowid

            for item in self.carrito:
                sub = item["precio"] * item["cantidad"]
                ganancia = (item["precio"] - item["costo"]) * item["cantidad"]
                row_stock = conn.execute(
                    "SELECT stock FROM productos WHERE id = ?",
                    (item["id"],)
                ).fetchone()
                if not row_stock:
                    raise RuntimeError(f'El producto "{item["nombre"]}" no existe en la base de datos.')
                stock_actual = float(row_stock[0])
                if stock_actual + 1e-9 < float(item["cantidad"]):
                    es_granel = bool(item.get("es_granel", False))
                    raise RuntimeError(
                        f'Stock insuficiente para "{item["nombre"]}". '
                        f'Disponible: {self._fmt_unidades(stock_actual, es_granel, es_granel)}'
                    )

                conn.execute(
                    "INSERT INTO detalle_venta"
                    " (venta_id,producto_id,nombre,precio,costo,cantidad,subtotal,ganancia,es_granel)"
                    " VALUES (?,?,?,?,?,?,?,?,?)",
                    (
                        venta_id, item["id"], item["nombre"], item["precio"], item["costo"],
                        item["cantidad"], sub, ganancia, int(bool(item.get("es_granel", False)))
                    )
                )
                conn.execute(
                    "UPDATE productos SET stock = stock - ? WHERE id = ?",
                    (item["cantidad"], item["id"])
                )

            if cliente_id is not None:
                conn.execute(
                    "UPDATE clientes"
                    " SET ultima_visita = ?, total_compras = IFNULL(total_compras,0) + ?"
                    " WHERE id = ?",
                    (fecha, total, cliente_id)
                )

            self.registrar_movimiento_venta(
                venta_id=venta_id,
                total=total,
                metodo_pago="EFECTIVO",
                conn=conn
            )

            conn.commit()
        except RuntimeError as e:
            if conn:
                conn.rollback()
            messagebox.showwarning("Venta no registrada", str(e), parent=self)
            return
        except sqlite3.Error as e:
            if conn:
                conn.rollback()
            messagebox.showerror(
                "Error de base de datos",
                f"No se pudo registrar la venta.\nDetalle técnico: {e}",
                parent=self
            )
            return
        except Exception as e:
            if conn:
                conn.rollback()
            messagebox.showerror(
                "Error inesperado",
                f"No se pudo completar la venta.\nDetalle: {e}",
                parent=self
            )
            return
        finally:
            if conn:
                conn.close()

        messagebox.showinfo("✔ Venta registrada",
            f"Venta #{venta_id} guardada.\nTotal: ${total:.2f}", parent=self)
        self.carrito.clear()
        self._refresh_carrito()
        self._cargar_productos()
        if hasattr(self, "tabla_clientes"):
            self._cargar_tabla_clientes()
        self._cargar_clientes_en_venta()
        if hasattr(self, "tabla_conta"):
            self._cargar_tabla_contabilidad()

    def _cliente_id_venta_actual(self):
        if not hasattr(self, "sv_cliente_venta"):
            return None
        opcion = self.sv_cliente_venta.get().strip()
        return self._cliente_id_por_opcion.get(opcion)

    def _cargar_clientes_cache(self):
        with get_conn() as conn:
            self._clientes_cache = conn.execute(
                "SELECT id,nombre,telefono,email,notas,ultima_visita,total_compras"
                " FROM clientes ORDER BY nombre"
            ).fetchall()

    def _cargar_clientes_en_venta(self):
        self._cargar_clientes_cache()
        if not hasattr(self, "cmb_cliente_venta"):
            return

        actual = self.sv_cliente_venta.get().strip() if hasattr(self, "sv_cliente_venta") else ""
        opciones = ["Público general"]
        self._cliente_id_por_opcion = {"Público general": None}

        for row in self._clientes_cache:
            cid, nombre, telefono = row[0], row[1], row[2]
            etiqueta = f"{nombre} · {telefono}" if telefono else nombre
            if etiqueta in self._cliente_id_por_opcion:
                etiqueta = f"{etiqueta} (ID {cid})"
            self._cliente_id_por_opcion[etiqueta] = cid
            opciones.append(etiqueta)

        self._cliente_opciones_venta = list(opciones)
        self.cmb_cliente_venta["values"] = tuple(opciones)
        if actual in opciones:
            self.sv_cliente_venta.set(actual)
        else:
            self.sv_cliente_venta.set("Público general")
        self._filtrar_clientes_en_venta()

    def _filtrar_clientes_en_venta(self):
        if not hasattr(self, "cmb_cliente_venta"):
            return
        opciones_base = list(getattr(self, "_cliente_opciones_venta", ["Público general"]))
        if not opciones_base:
            opciones_base = ["Público general"]

        q = ""
        if hasattr(self, "sv_cliente_busqueda_venta"):
            q = self._texto_busqueda(self.sv_cliente_busqueda_venta.get())

        actual = self.sv_cliente_venta.get().strip() if hasattr(self, "sv_cliente_venta") else ""
        if actual not in self._cliente_id_por_opcion:
            self.sv_cliente_venta.set("Público general")

        if not hasattr(self, "lst_clientes_venta") or not hasattr(self, "frame_lista_clientes_venta"):
            return

        self.lst_clientes_venta.delete(0, "end")
        if not q:
            self.frame_lista_clientes_venta.pack_forget()
            return

        filtradas = [op for op in opciones_base if q in self._texto_busqueda(op)]
        if not filtradas:
            filtradas = ["Sin coincidencias"]

        for op in filtradas:
            self.lst_clientes_venta.insert("end", op)
        self.lst_clientes_venta.config(height=min(max(len(filtradas), 1), 6))
        self.frame_lista_clientes_venta.pack(fill="x", pady=(0, 4))

    def _seleccionar_cliente_lista_venta(self, event=None):
        if not hasattr(self, "lst_clientes_venta"):
            return
        sel = self.lst_clientes_venta.curselection()
        if not sel:
            return
        etiqueta = self.lst_clientes_venta.get(sel[0])
        if etiqueta not in self._cliente_id_por_opcion:
            return
        self.sv_cliente_venta.set(etiqueta)
        if hasattr(self, "sv_cliente_busqueda_venta"):
            self.sv_cliente_busqueda_venta.set("")

    # ── Contabilidad (núcleo) ─────────────────────────────
    def _cargar_categorias_contables_cache(self):
        with get_conn() as conn:
            self._categorias_contables_cache = conn.execute(
                "SELECT nombre,tipo,naturaleza"
                " FROM categorias_contables"
                " WHERE activa = 1"
                " ORDER BY tipo, nombre"
            ).fetchall()

    def _categorias_por_tipo(self, tipo):
        return [c[0] for c in self._categorias_contables_cache if c[1] == tipo]

    def _opciones_subtipo_por_tipo(self, tipo):
        if tipo == "INGRESO":
            return ["VENTA", "INGRESO_EXTRA", "DEPOSITO", "AJUSTE"]
        return ["GASTO", "RETIRO_DUENO", "AJUSTE"]

    def _categoria_default(self, tipo, subtipo):
        if tipo == "INGRESO" and subtipo == "VENTA":
            return "Ventas"
        if tipo == "INGRESO" and subtipo == "INGRESO_EXTRA":
            return "Ingreso extra"
        if tipo == "INGRESO" and subtipo == "DEPOSITO":
            return "Depósito"
        if tipo == "INGRESO" and subtipo == "AJUSTE":
            return "Ajuste de caja (+)"
        if tipo == "EGRESO" and subtipo == "RETIRO_DUENO":
            return "Retiro dueño"
        if tipo == "EGRESO" and subtipo == "AJUSTE":
            return "Ajuste de caja (-)"
        if tipo == "EGRESO" and subtipo == "GASTO":
            for cand in ("Compra inventario", "Servicios", "Transporte", "Mantenimiento", "Renta", "Sueldos"):
                if cand in self._categorias_por_tipo("EGRESO"):
                    return cand
        cats = self._categorias_por_tipo(tipo)
        return cats[0] if cats else ""

    def _registrar_movimiento_caja(
        self, tipo, subtipo, categoria, concepto, monto,
        metodo_pago="EFECTIVO", referencia="", venta_id=None, usuario="admin", notas="", conn=None
    ):
        tipo = (tipo or "").strip().upper()
        subtipo = (subtipo or "").strip().upper()
        if tipo not in ("INGRESO", "EGRESO"):
            raise ValueError("Tipo de movimiento inválido.")
        if subtipo not in ("VENTA", "GASTO", "RETIRO_DUENO", "DEPOSITO", "AJUSTE", "INGRESO_EXTRA"):
            raise ValueError("Subtipo de movimiento inválido.")
        if not categoria or not categoria.strip():
            raise ValueError("La categoría es obligatoria.")
        if not concepto or not concepto.strip():
            raise ValueError("El concepto es obligatorio.")
        monto = float(monto)
        if monto <= 0:
            raise ValueError("El monto debe ser mayor a 0.")

        def _validar_categoria(conn_obj):
            cat = conn_obj.execute(
                "SELECT tipo FROM categorias_contables WHERE nombre = ? AND activa = 1",
                (categoria.strip(),)
            ).fetchone()
            if not cat:
                raise ValueError(f'La categoría "{categoria}" no existe o está inactiva.')
            tipo_cat = str(cat[0]).strip().upper()
            if tipo_cat != tipo:
                raise ValueError(
                    f'La categoría "{categoria}" es de tipo {tipo_cat}, no {tipo}.'
                )

        sql = (
            "INSERT INTO movimientos_caja"
            " (fecha,tipo,subtipo,categoria,concepto,metodo_pago,monto,referencia,venta_id,usuario,notas)"
            " VALUES (datetime('now','localtime'),?,?,?,?,?,?,?,?,?,?)"
        )
        vals = (
            tipo, subtipo, categoria.strip(), concepto.strip(), (metodo_pago or "EFECTIVO").strip().upper(),
            monto, (referencia or "").strip(), venta_id, (usuario or "admin").strip(), (notas or "").strip()
        )

        if conn is not None:
            _validar_categoria(conn)
            cur = conn.execute(sql, vals)
            return cur.lastrowid
        with get_conn() as own:
            _validar_categoria(own)
            cur = own.execute(sql, vals)
            return cur.lastrowid

    def registrar_movimiento_venta(self, venta_id, total, metodo_pago="EFECTIVO", conn=None):
        if float(total) <= 0:
            return None
        return self._registrar_movimiento_caja(
            tipo="INGRESO",
            subtipo="VENTA",
            categoria="Ventas",
            concepto=f"Venta #{venta_id}",
            monto=total,
            metodo_pago=metodo_pago,
            referencia=str(venta_id),
            venta_id=venta_id,
            conn=conn
        )

    def sincronizar_venta_en_caja(self, venta_id):
        with get_conn() as conn:
            venta = conn.execute(
                "SELECT id,total FROM ventas WHERE id = ?",
                (venta_id,)
            ).fetchone()
            conn.execute("DELETE FROM movimientos_caja WHERE venta_id = ?", (venta_id,))
            if venta:
                self.registrar_movimiento_venta(venta[0], venta[1], conn=conn)

    def registrar_gasto(self, categoria, concepto, monto, naturaleza=None, referencia=""):
        return self._registrar_movimiento_caja(
            tipo="EGRESO",
            subtipo="GASTO",
            categoria=categoria,
            concepto=concepto,
            monto=monto,
            metodo_pago="EFECTIVO",
            referencia=referencia
        )

    def registrar_retiro_dueno(self, monto, concepto="Retiro de dueño"):
        return self._registrar_movimiento_caja(
            tipo="EGRESO",
            subtipo="RETIRO_DUENO",
            categoria="Retiro dueño",
            concepto=concepto,
            monto=monto,
            metodo_pago="EFECTIVO"
        )

    def registrar_ingreso_extra(self, categoria, concepto, monto):
        return self._registrar_movimiento_caja(
            tipo="INGRESO",
            subtipo="INGRESO_EXTRA",
            categoria=categoria,
            concepto=concepto,
            monto=monto,
            metodo_pago="EFECTIVO"
        )

    def abrir_caja(self, saldo_inicial, usuario="admin"):
        saldo_inicial = float(saldo_inicial)
        if saldo_inicial < 0:
            raise ValueError("El saldo inicial no puede ser negativo.")
        with get_conn() as conn:
            abierta = conn.execute(
                "SELECT id FROM cortes_caja WHERE estado='ABIERTO' ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if abierta:
                raise RuntimeError("Ya existe una caja abierta.")
            cur = conn.execute(
                "INSERT INTO cortes_caja (fecha_apertura,saldo_inicial,usuario_apertura)"
                " VALUES (datetime('now','localtime'),?,?)",
                (saldo_inicial, usuario)
            )
            return cur.lastrowid

    def resumen_caja_actual(self):
        with get_conn() as conn:
            corte = conn.execute(
                "SELECT id,fecha_apertura,saldo_inicial FROM cortes_caja"
                " WHERE estado='ABIERTO' ORDER BY id DESC LIMIT 1"
            ).fetchone()
            if not corte:
                return None
            corte_id, fecha_ap, saldo_inicial = corte
            suma = conn.execute(
                "SELECT"
                " IFNULL(SUM(CASE WHEN tipo='INGRESO' AND subtipo='VENTA' AND metodo_pago='EFECTIVO' THEN monto END),0),"
                " IFNULL(SUM(CASE WHEN tipo='INGRESO' AND subtipo<>'VENTA' THEN monto END),0),"
                " IFNULL(SUM(CASE WHEN tipo='EGRESO' THEN monto END),0)"
                " FROM movimientos_caja"
                " WHERE fecha >= ?",
                (fecha_ap,)
            ).fetchone()
            ventas_efectivo, ingresos_extra, egresos = [float(x or 0) for x in suma]
            saldo_teorico = float(saldo_inicial) + ventas_efectivo + ingresos_extra - egresos
            conn.execute(
                "UPDATE cortes_caja"
                " SET ventas_efectivo=?, ingresos_extra=?, egresos=?, saldo_teorico=?"
                " WHERE id=?",
                (ventas_efectivo, ingresos_extra, egresos, saldo_teorico, corte_id)
            )
            return {
                "id": corte_id,
                "fecha_apertura": fecha_ap,
                "saldo_inicial": float(saldo_inicial),
                "ventas_efectivo": ventas_efectivo,
                "ingresos_extra": ingresos_extra,
                "egresos": egresos,
                "saldo_teorico": saldo_teorico,
            }

    def cerrar_caja(self, saldo_real, usuario="admin"):
        saldo_real = float(saldo_real)
        resumen = self.resumen_caja_actual()
        if not resumen:
            raise RuntimeError("No hay caja abierta para cerrar.")
        diferencia = saldo_real - resumen["saldo_teorico"]
        with get_conn() as conn:
            conn.execute(
                "UPDATE cortes_caja"
                " SET fecha_cierre=datetime('now','localtime'),"
                " saldo_real=?, diferencia=?, estado='CERRADO', usuario_cierre=?"
                " WHERE id=?",
                (saldo_real, diferencia, usuario, resumen["id"])
            )
        return diferencia

    def exportar_libro_diario_csv(self, fecha_ini, fecha_fin, ruta):
        inicio = f"{fecha_ini} 00:00:00"
        fin = f"{fecha_fin} 23:59:59"
        with get_conn() as conn:
            rows = conn.execute(
                "SELECT m.fecha, m.id, m.tipo, m.subtipo, m.categoria,"
                " IFNULL(c.naturaleza,'NO_APLICA') AS naturaleza,"
                " m.concepto, m.metodo_pago, m.monto, m.venta_id, m.referencia, m.usuario"
                " FROM movimientos_caja m"
                " LEFT JOIN categorias_contables c ON c.nombre = m.categoria"
                " WHERE m.fecha >= ? AND m.fecha <= ?"
                " ORDER BY m.id ASC",
                (inicio, fin)
            ).fetchall()
        with open(ruta, mode="w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow([
                "fecha", "folio", "tipo", "subtipo", "categoria", "naturaleza",
                "concepto", "metodo_pago", "monto", "venta_id", "referencia", "usuario"
            ])
            for r in rows:
                rr = list(r)
                rr[8] = f"{float(rr[8]):.2f}"  # monto
                w.writerow(rr)
        return len(rows)

    def reporte_resultado(self, fecha_ini, fecha_fin):
        inicio = f"{fecha_ini} 00:00:00"
        fin = f"{fecha_fin} 23:59:59"
        with get_conn() as conn:
            ingresos, egresos = conn.execute(
                "SELECT"
                " IFNULL(SUM(CASE WHEN tipo='INGRESO' THEN monto END),0),"
                " IFNULL(SUM(CASE WHEN tipo='EGRESO' THEN monto END),0)"
                " FROM movimientos_caja"
                " WHERE fecha >= ? AND fecha <= ?",
                (inicio, fin)
            ).fetchone()
            gastos_operativos = conn.execute(
                "SELECT IFNULL(SUM(monto),0)"
                " FROM movimientos_caja"
                " WHERE tipo='EGRESO' AND subtipo='GASTO'"
                " AND fecha >= ? AND fecha <= ?",
                (inicio, fin)
            ).fetchone()[0]

        ingresos = float(ingresos or 0)
        egresos = float(egresos or 0)
        gastos_operativos = float(gastos_operativos or 0)
        utilidad_operativa = ingresos - gastos_operativos
        utilidad_neta_estimada = ingresos - egresos
        margen_operativo = (utilidad_operativa / ingresos) if ingresos > 0 else 0
        return {
            "ingresos": ingresos,
            "egresos": egresos,
            "gastos_operativos": gastos_operativos,
            "utilidad_operativa": utilidad_operativa,
            "utilidad_neta_estimada": utilidad_neta_estimada,
            "margen_operativo": margen_operativo,
        }

    def reporte_gastos_por_categoria(self, fecha_ini, fecha_fin):
        inicio = f"{fecha_ini} 00:00:00"
        fin = f"{fecha_fin} 23:59:59"
        with get_conn() as conn:
            ventas = conn.execute(
                "SELECT IFNULL(SUM(monto),0)"
                " FROM movimientos_caja"
                " WHERE tipo='INGRESO' AND subtipo='VENTA'"
                " AND fecha >= ? AND fecha <= ?",
                (inicio, fin)
            ).fetchone()[0]
            rows = conn.execute(
                "SELECT categoria, IFNULL(SUM(monto),0) AS total"
                " FROM movimientos_caja"
                " WHERE tipo='EGRESO' AND subtipo='GASTO'"
                " AND fecha >= ? AND fecha <= ?"
                " GROUP BY categoria"
                " ORDER BY total DESC",
                (inicio, fin)
            ).fetchall()
        ventas = float(ventas or 0)
        out = []
        for categoria, total in rows:
            total = float(total or 0)
            pct = (total / ventas * 100) if ventas > 0 else 0
            out.append({"categoria": categoria, "total": total, "porcentaje_sobre_ventas": pct})
        return out

    def reporte_fijos_vs_variables(self, mes):
        # mes esperado: 'YYYY-MM'
        inicio = f"{mes}-01 00:00:00"
        y, m = [int(x) for x in mes.split("-")]
        if m == 12:
            next_month = datetime.date(y + 1, 1, 1)
        else:
            next_month = datetime.date(y, m + 1, 1)
        fin_date = next_month - datetime.timedelta(days=1)
        fin = f"{fin_date.isoformat()} 23:59:59"
        with get_conn() as conn:
            fijo, variable = conn.execute(
                "SELECT"
                " IFNULL(SUM(CASE WHEN c.naturaleza='FIJO' THEN m.monto END),0),"
                " IFNULL(SUM(CASE WHEN c.naturaleza='VARIABLE' THEN m.monto END),0)"
                " FROM movimientos_caja m"
                " LEFT JOIN categorias_contables c ON c.nombre = m.categoria"
                " WHERE m.tipo='EGRESO' AND m.subtipo='GASTO'"
                " AND m.fecha >= ? AND m.fecha <= ?",
                (inicio, fin)
            ).fetchone()
        fijo = float(fijo or 0)
        variable = float(variable or 0)
        return {"fijo": fijo, "variable": variable, "total": fijo + variable}

    def flujo_efectivo_diario(self, fecha_ini, fecha_fin):
        inicio = f"{fecha_ini} 00:00:00"
        fin = f"{fecha_fin} 23:59:59"
        with get_conn() as conn:
            rows = conn.execute(
                "SELECT date(fecha) AS dia,"
                " IFNULL(SUM(CASE WHEN tipo='INGRESO' THEN monto END),0) AS ingresos,"
                " IFNULL(SUM(CASE WHEN tipo='EGRESO' THEN monto END),0) AS egresos"
                " FROM movimientos_caja"
                " WHERE fecha >= ? AND fecha <= ?"
                " GROUP BY date(fecha)"
                " ORDER BY dia ASC",
                (inicio, fin)
            ).fetchall()
        return [
            {
                "fecha": r[0],
                "ingresos": float(r[1] or 0),
                "egresos": float(r[2] or 0),
                "flujo_neto": float(r[1] or 0) - float(r[2] or 0),
            }
            for r in rows
        ]

    def exportar_resumen_mensual_csv(self, anio, mes, ruta):
        mes_txt = f"{int(anio):04d}-{int(mes):02d}"
        inicio = f"{mes_txt}-01 00:00:00"
        y, m = int(anio), int(mes)
        if m == 12:
            next_month = datetime.date(y + 1, 1, 1)
        else:
            next_month = datetime.date(y, m + 1, 1)
        fin_date = next_month - datetime.timedelta(days=1)
        fin = f"{fin_date.isoformat()} 23:59:59"

        with get_conn() as conn:
            rows = conn.execute(
                "SELECT m.tipo, m.categoria, IFNULL(c.naturaleza,'NO_APLICA') AS naturaleza,"
                " IFNULL(SUM(m.monto),0) AS total"
                " FROM movimientos_caja m"
                " LEFT JOIN categorias_contables c ON c.nombre = m.categoria"
                " WHERE m.fecha >= ? AND m.fecha <= ?"
                " GROUP BY m.tipo, m.categoria, naturaleza"
                " ORDER BY m.tipo, total DESC",
                (inicio, fin)
            ).fetchall()

        with open(ruta, mode="w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(["tipo", "categoria", "naturaleza", "total"])
            for tipo, categoria, naturaleza, total in rows:
                w.writerow([tipo, categoria, naturaleza, f"{float(total):.2f}"])
        return len(rows)

    def exportar_cortes_caja_csv(self, fecha_ini, fecha_fin, ruta):
        inicio = f"{fecha_ini} 00:00:00"
        fin = f"{fecha_fin} 23:59:59"
        with get_conn() as conn:
            rows = conn.execute(
                "SELECT id,fecha_apertura,fecha_cierre,saldo_inicial,ventas_efectivo,ingresos_extra,"
                " egresos,saldo_teorico,saldo_real,diferencia,estado,usuario_apertura,usuario_cierre"
                " FROM cortes_caja"
                " WHERE fecha_apertura >= ? AND fecha_apertura <= ?"
                " ORDER BY id DESC",
                (inicio, fin)
            ).fetchall()
        with open(ruta, mode="w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow([
                "id", "fecha_apertura", "fecha_cierre", "saldo_inicial", "ventas_efectivo",
                "ingresos_extra", "egresos", "saldo_teorico", "saldo_real", "diferencia",
                "estado", "usuario_apertura", "usuario_cierre"
            ])
            for r in rows:
                rr = list(r)
                for idx in (3, 4, 5, 6, 7, 8, 9):
                    if rr[idx] is not None:
                        rr[idx] = f"{float(rr[idx]):.2f}"
                w.writerow(rr)
        return len(rows)

    def _set_caducidad_form(self, valor):
        fecha = (valor or "").strip()
        if hasattr(self, "sv_caducidad"):
            self.sv_caducidad.set(fecha)
            return
        if "e_caducidad" in getattr(self, "_prod_entries", {}):
            ent = self._prod_entries["e_caducidad"]
            ent.delete(0, "end")
            ent.insert(0, fecha)

    def _get_caducidad_form(self):
        if hasattr(self, "sv_caducidad"):
            return self.sv_caducidad.get().strip()
        if "e_caducidad" in getattr(self, "_prod_entries", {}):
            return self._prod_entries["e_caducidad"].get().strip()
        return ""

    def _normalizar_categoria_producto(self, texto):
        return " ".join((texto or "").strip().split())

    def _cargar_categorias_producto_cache(self):
        categorias = set()
        with get_conn() as conn:
            rows_catalogo = conn.execute(
                "SELECT nombre FROM categorias_productos ORDER BY nombre COLLATE NOCASE"
            ).fetchall()
            rows_productos = conn.execute(
                "SELECT DISTINCT categoria FROM productos"
                " WHERE categoria IS NOT NULL AND TRIM(categoria) <> ''"
            ).fetchall()
        for (nombre,) in rows_catalogo + rows_productos:
            cat = self._normalizar_categoria_producto(nombre)
            if cat:
                categorias.add(cat)
        if "General" not in categorias:
            categorias.add("General")
        self._categorias_producto_cache = sorted(categorias, key=str.casefold)

    def _actualizar_combo_categoria_producto(self, filtro=None, abrir=False):
        if not hasattr(self, "cmb_categoria_producto"):
            return
        termino = self._normalizar_categoria_producto(
            self.sv_categoria_producto.get() if filtro is None else filtro
        )
        if termino:
            fold = termino.casefold()
            opciones = [c for c in self._categorias_producto_cache if fold in c.casefold()]
        else:
            opciones = list(self._categorias_producto_cache)
        self.cmb_categoria_producto["values"] = tuple(opciones)
        if abrir and opciones:
            self.after_idle(lambda: self.cmb_categoria_producto.event_generate("<Down>"))

    def _refrescar_categorias_producto(self, filtro=None):
        self._cargar_categorias_producto_cache()
        self._actualizar_combo_categoria_producto(filtro=filtro)

    def _on_categoria_producto_keyrelease(self, event=None):
        if event and event.keysym in {
            "Up", "Down", "Left", "Right", "Home", "End", "Prior", "Next",
            "Return", "Escape", "Tab", "Shift_L", "Shift_R", "Control_L",
            "Control_R", "Alt_L", "Alt_R"
        }:
            return
        self._actualizar_combo_categoria_producto(abrir=True)

    def _registrar_categoria_producto(self, categoria):
        nombre = self._normalizar_categoria_producto(categoria) or "General"
        with get_conn() as conn:
            row = conn.execute(
                "SELECT nombre FROM categorias_productos WHERE LOWER(nombre)=LOWER(?)",
                (nombre,)
            ).fetchone()
            if row:
                return row[0]
            conn.execute(
                "INSERT INTO categorias_productos (nombre) VALUES (?)",
                (nombre,)
            )
        return nombre

    def _crear_categoria_producto(self):
        nombre = self._normalizar_categoria_producto(self.sv_categoria_producto.get())
        if not nombre:
            messagebox.showwarning(
                "Categoría vacía",
                "Escribe un nombre de categoría para crearla.",
                parent=self
            )
            return
        with get_conn() as conn:
            row = conn.execute(
                "SELECT nombre FROM categorias_productos WHERE LOWER(nombre)=LOWER(?)",
                (nombre,)
            ).fetchone()
            if row:
                final = row[0]
            else:
                conn.execute(
                    "INSERT INTO categorias_productos (nombre) VALUES (?)",
                    (nombre,)
                )
                final = nombre
        self.sv_categoria_producto.set(final)
        self._refrescar_categorias_producto(filtro=final)
        if row:
            messagebox.showinfo(
                "Categoría existente",
                f'La categoría "{final}" ya estaba registrada.',
                parent=self
            )
        else:
            messagebox.showinfo(
                "Categoría creada",
                f'Se agregó la categoría "{final}".',
                parent=self
            )

    def _set_prod_field(self, key, valor):
        if key == "e_caducidad":
            self._set_caducidad_form(str(valor))
            return
        ent = self._prod_entries[key]
        ent.delete(0, "end")
        ent.insert(0, str(valor))

    def _clear_prod_form_inputs(self):
        for key, ent in self._prod_entries.items():
            if key == "e_caducidad":
                self._set_caducidad_form("")
            else:
                ent.delete(0, "end")

    def _abrir_selector_fecha(self):
        if not hasattr(self, "sv_caducidad"):
            return

        hoy = datetime.date.today()
        base = hoy
        actual = self._get_caducidad_form()
        if actual:
            try:
                base = datetime.date.fromisoformat(actual)
            except ValueError:
                pass

        dlg = tk.Toplevel(self)
        dlg.title("Seleccionar fecha de caducidad")
        dlg.configure(bg=C["card"])
        dlg.resizable(False, False)
        dlg.transient(self)
        dlg.grab_set()
        dlg.focus_set()
        self.update_idletasks()
        x = self.winfo_x() + (self.winfo_width()  // 2) - 220
        y = self.winfo_y() + (self.winfo_height() // 2) - 130
        dlg.geometry(f"440x260+{x}+{y}")

        tk.Label(dlg, text="FECHA DE CADUCIDAD", fg=C["accent2"], bg=C["card"],
                 font=("Courier", 11, "bold")).pack(pady=(16, 6))
        tk.Label(dlg, text="Selecciona año, mes y día.", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).pack()

        frame = tk.Frame(dlg, bg=C["card"])
        frame.pack(pady=14)

        years = [str(y) for y in range(hoy.year - 5, hoy.year + 11)]
        if str(base.year) not in years:
            years.append(str(base.year))
            years.sort()
        months = [f"{m:02d}" for m in range(1, 13)]

        sv_y = tk.StringVar(value=str(base.year))
        sv_m = tk.StringVar(value=f"{base.month:02d}")
        sv_d = tk.StringVar(value=f"{base.day:02d}")

        def _dias_del_mes(anio, mes):
            inicio = datetime.date(anio, mes, 1)
            if mes == 12:
                prox = datetime.date(anio + 1, 1, 1)
            else:
                prox = datetime.date(anio, mes + 1, 1)
            return (prox - inicio).days

        tk.Label(frame, text="Año", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).grid(row=0, column=0, padx=4, sticky="w")
        tk.Label(frame, text="Mes", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).grid(row=0, column=1, padx=4, sticky="w")
        tk.Label(frame, text="Día", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).grid(row=0, column=2, padx=4, sticky="w")

        cb_y = ttk.Combobox(frame, textvariable=sv_y, values=years, state="readonly", width=8)
        cb_m = ttk.Combobox(frame, textvariable=sv_m, values=months, state="readonly", width=6)
        cb_d = ttk.Combobox(frame, textvariable=sv_d, state="readonly", width=6)
        cb_y.grid(row=1, column=0, padx=4, ipady=2)
        cb_m.grid(row=1, column=1, padx=4, ipady=2)
        cb_d.grid(row=1, column=2, padx=4, ipady=2)

        def _actualizar_dias(*_):
            try:
                anio = int(sv_y.get())
                mes = int(sv_m.get())
            except ValueError:
                return
            max_dia = _dias_del_mes(anio, mes)
            dias = [f"{d:02d}" for d in range(1, max_dia + 1)]
            cb_d.config(values=dias)
            if sv_d.get() not in dias:
                sv_d.set(dias[-1])

        _actualizar_dias()
        cb_y.bind("<<ComboboxSelected>>", _actualizar_dias)
        cb_m.bind("<<ComboboxSelected>>", _actualizar_dias)

        lbl_error = tk.Label(dlg, text="", fg=C["red"], bg=C["card"],
                             font=("Courier", 9))
        lbl_error.pack(pady=(2, 6))

        def _aceptar():
            try:
                fecha = datetime.date(int(sv_y.get()), int(sv_m.get()), int(sv_d.get()))
            except ValueError:
                lbl_error.config(text="Fecha inválida.")
                return
            self._set_caducidad_form(fecha.isoformat())
            dlg.destroy()

        def _limpiar():
            self._set_caducidad_form("")
            dlg.destroy()

        btns = tk.Frame(dlg, bg=C["card"])
        btns.pack(fill="x", padx=18, pady=(2, 12))
        tk.Button(btns, text="🗑 Limpiar fecha", bg=C["panel"], fg=C["muted"], bd=0,
                  font=("Courier", 10), pady=8, cursor="hand2",
                  command=_limpiar).pack(side="left", padx=(0, 4))
        tk.Button(btns, text="✕ Cancelar", bg=C["panel"], fg=C["muted"], bd=0,
                  font=("Courier", 10), pady=8, cursor="hand2",
                  command=dlg.destroy).pack(side="right", padx=(4, 0))
        tk.Button(btns, text="✔ Aceptar", bg=C["green"], fg=C["white"], bd=0,
                  font=("Courier", 10, "bold"), pady=8, cursor="hand2",
                  activebackground="#27ae60",
                  command=_aceptar).pack(side="right")

        dlg.bind("<Return>", lambda e: _aceptar())
        dlg.bind("<Escape>", lambda e: dlg.destroy())
        dlg.protocol("WM_DELETE_WINDOW", dlg.destroy)

    # ══════════════════════════════════════════════════════
    #  PÁGINA: PRODUCTOS
    # ══════════════════════════════════════════════════════
    def _build_page_productos(self):
        page = tk.Frame(self.content, bg=C["bg"])
        self.pages["productos"] = page

        form_card = tk.Frame(page, bg=C["card"], padx=16, pady=14)
        form_card.pack(fill="x", pady=(0,12))

        tk.Label(form_card, text="AGREGAR / EDITAR PRODUCTO",
                 fg=C["muted"], bg=C["card"], font=("Courier", 9, "bold")).grid(
                     row=0, column=0, columnspan=7, sticky="w", pady=(0,10))

        fields = [("Código", "e_codigo"), ("Nombre", "e_nombre"),
                  ("Costo $", "e_costo"), ("Precio venta $", "e_precio"),
                  ("Stock", "e_stock"), ("Categoría", "e_categoria"),
                  ("Caducidad (YYYY-MM-DD)", "e_caducidad")]
        entry_widths = {
            "e_codigo": 12,
            "e_nombre": 30,
            "e_costo": 8,
            "e_precio": 10,
            "e_stock": 7,
            "e_categoria": 16,
            "e_caducidad": 14,
        }
        self._prod_entries = {}
        self.sv_caducidad = tk.StringVar(value="")
        self.sv_categoria_producto = tk.StringVar(value="")
        self._refrescar_categorias_producto()
        for col, (lbl, key) in enumerate(fields):
            color_lbl = C["yellow"] if key == "e_costo" else C["muted"]
            tk.Label(form_card, text=lbl, fg=color_lbl, bg=C["card"],
                     font=("Courier", 9)).grid(row=1, column=col, padx=(0,4), sticky="w")

            if key == "e_codigo":
                wrap = tk.Frame(form_card, bg=C["card"])
                wrap.grid(row=2, column=col, padx=(0,8), sticky="ew")
                ent = tk.Entry(wrap, bg=C["panel"], fg=C["text"],
                               insertbackground=C["text"], bd=0, font=("Courier", 11),
                               highlightbackground=C["border"], highlightthickness=1,
                               width=entry_widths[key])
                ent.pack(side="left", fill="x", expand=True, ipady=6)
                tk.Button(wrap, text="⟳", bg=C["accent2"], fg=C["white"],
                          bd=0, font=("Courier", 10), padx=5, cursor="hand2",
                          activebackground="#6a4aaf",
                          command=self._limpiar_form_producto).pack(side="left", padx=(2,0))
            elif key == "e_caducidad":
                wrap = tk.Frame(form_card, bg=C["card"])
                wrap.grid(row=2, column=col, padx=(0,8), sticky="ew")
                ent = tk.Entry(
                    wrap,
                    textvariable=self.sv_caducidad,
                    bg=C["panel"], fg=C["text"],
                    readonlybackground=C["panel"],
                    insertbackground=C["text"],
                    bd=0, font=("Courier", 11),
                    highlightbackground=C["border"], highlightthickness=1,
                    width=entry_widths[key], state="readonly", cursor="hand2"
                )
                ent.pack(side="left", fill="x", expand=True, ipady=6)
                tk.Button(
                    wrap, text="📅", bg=C["accent2"], fg=C["white"], bd=0,
                    font=("Courier", 10), padx=5, cursor="hand2",
                    activebackground="#6a4aaf",
                    command=self._abrir_selector_fecha
                ).pack(side="left", padx=(2,0))
            elif key == "e_categoria":
                wrap = tk.Frame(form_card, bg=C["card"])
                wrap.grid(row=2, column=col, padx=(0,8), sticky="ew")
                ent = ttk.Combobox(
                    wrap,
                    textvariable=self.sv_categoria_producto,
                    state="normal",
                    values=tuple(self._categorias_producto_cache),
                    width=entry_widths[key],
                    font=("Courier", 10)
                )
                ent.pack(side="left", fill="x", expand=True, ipady=3)
                ent.bind("<KeyRelease>", self._on_categoria_producto_keyrelease)
                ent.bind("<FocusIn>", lambda _e: self._actualizar_combo_categoria_producto())
                ent.bind("<<ComboboxSelected>>", lambda _e: self._actualizar_combo_categoria_producto())
                self.cmb_categoria_producto = ent
                tk.Button(
                    wrap, text="＋", bg=C["accent2"], fg=C["white"], bd=0,
                    font=("Courier", 10), padx=5, cursor="hand2",
                    activebackground="#6a4aaf",
                    command=self._crear_categoria_producto
                ).pack(side="left", padx=(2,0))
            else:
                ent = tk.Entry(form_card, bg=C["panel"], fg=C["text"],
                               insertbackground=C["text"], bd=0, font=("Courier", 11),
                               highlightbackground=C["border"], highlightthickness=1,
                               width=entry_widths[key])
                ent.grid(row=2, column=col, padx=(0,8), ipady=6, sticky="ew")

            self._prod_entries[key] = ent
        form_card.columnconfigure(0, weight=2, minsize=120)   # Código
        form_card.columnconfigure(1, weight=7, minsize=300)   # Nombre
        form_card.columnconfigure(2, weight=2, minsize=95)    # Costo
        form_card.columnconfigure(3, weight=2, minsize=110)   # Precio
        form_card.columnconfigure(4, weight=1, minsize=85)    # Stock
        form_card.columnconfigure(5, weight=3, minsize=160)   # Categoría
        form_card.columnconfigure(6, weight=2, minsize=150)   # Caducidad
        self.sv_es_granel = tk.BooleanVar(value=False)
        tk.Checkbutton(
            form_card,
            text="¿Producto a granel? (stock decimal en kg)",
            variable=self.sv_es_granel,
            onvalue=True,
            offvalue=False,
            bg=C["card"],
            fg=C["text"],
            selectcolor=C["panel"],
            activebackground=C["card"],
            activeforeground=C["text"],
            font=("Courier", 9),
            highlightthickness=0,
            bd=0
        ).grid(row=3, column=0, columnspan=7, sticky="w", pady=(10, 0))

        btn_frame = tk.Frame(form_card, bg=C["card"])
        btn_frame.grid(row=4, column=0, columnspan=7, pady=(10,0), sticky="e")
        tk.Button(btn_frame, text="● Producto nuevo", bg=C["accent2"], fg=C["white"],
              bd=0, font=("Courier", 10), padx=12, pady=6, cursor="hand2",
              activebackground="#6a4aaf",
              command=self._nuevo_producto).pack(side="left", padx=(0,4))

        tk.Button(btn_frame, text="＋ Guardar", bg=C["accent"], fg=C["white"],
              bd=0, font=("Courier", 10, "bold"), padx=12, pady=6, cursor="hand2",
              command=self._guardar_producto).pack(side="left", padx=(0,4))
        tk.Button(btn_frame, text="✕ Eliminar", bg=C["red"], fg=C["white"],
              bd=0, font=("Courier", 10), padx=12, pady=6, cursor="hand2",
              command=self._eliminar_producto).pack(side="left", padx=(0,4))
        tk.Button(btn_frame, text="📋 Inventario Físico", bg=C["green"], fg=C["white"],
              bd=0, font=("Courier", 10, "bold"), padx=12, pady=6, cursor="hand2",
              activebackground="#1a4a10",
              command=self._exportar_inventario_docx).pack(side="left")

        cols = ("id","codigo","nombre","costo","precio","stock","categoria","caducidad","granel","estado")
        heads = ("ID","Código","Nombre","Costo","Precio venta","Stock","Categoría","Caducidad","A granel","Estado")
        widths = [40,75,180,70,80,70,90,120,70,110]

        frame_t = tk.Frame(page, bg=C["bg"])
        frame_t.pack(fill="both", expand=True)

        self.tabla_prod = ttk.Treeview(frame_t, columns=cols, show="headings",
                                       style="POS.Treeview", selectmode="browse")
        for c,h,w in zip(cols,heads,widths):
            self.tabla_prod.heading(c, text=h, anchor="center")
            self.tabla_prod.column(c, width=w, anchor="center", stretch=True)
        sb = ttk.Scrollbar(frame_t, orient="vertical",
                   command=self.tabla_prod.yview)
        self.tabla_prod.configure(yscrollcommand=sb.set)
        self.tabla_prod.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.tabla_prod.bind("<<TreeviewSelect>>", self._llenar_form_producto)

        search_f = tk.Frame(page, bg=C["bg"])
        search_f.pack(fill="x", pady=(0,6))
        tk.Label(search_f, text="Filtrar: ", fg=C["muted"], bg=C["bg"],
                 font=("Courier",10)).pack(side="left")
        self.sv_prod_filter = tk.StringVar()
        self.sv_prod_filter.trace_add("write", lambda *a: self._cargar_tabla_productos())
        tk.Entry(search_f, textvariable=self.sv_prod_filter,
                 bg=C["panel"], fg=C["text"], insertbackground=C["text"],
                 bd=0, font=("Courier",10), highlightthickness=1,
                 highlightbackground=C["border"]).pack(
                     side="left", fill="x", expand=True, padx=(6,0), ipady=5
                 )

        search_f.pack_forget()
        form_card.pack_forget()
        frame_t.pack_forget()
        form_card.pack(fill="x", pady=(0,8))
        search_f.pack(fill="x", pady=(0,6))
        frame_t.pack(fill="both", expand=True)

    def _cargar_tabla_productos(self):
        q = ""
        if hasattr(self, "sv_prod_filter"):
            q = self._texto_busqueda(self.sv_prod_filter.get())
        for row in self.tabla_prod.get_children():
            self.tabla_prod.delete(row)
        with get_conn() as conn:
            rows = conn.execute(
                "SELECT id,codigo,nombre,costo,precio,stock,categoria,a_granel,caducidad"
                " FROM productos ORDER BY nombre"
            ).fetchall()
        rows = sorted(rows, key=lambda r: self._clave_prioridad_salida(r[8], r[2]))
        for r in rows:
            codigo_q = self._texto_busqueda(r[1])
            nombre_q = self._texto_busqueda(r[2])
            categoria_q = self._texto_busqueda(r[6])
            if q in codigo_q or q in nombre_q or q in categoria_q:
                es_critico, dias = self._estado_caducidad(r[8])
                tags = []
                if r[5] <= 5:
                    tags.append("low")
                if es_critico:
                    tags.append("critical")
                stock_txt = self._fmt_unidades(r[5], bool(r[7]), bool(r[7]))
                estado = self._texto_estado_caducidad(dias)
                nombre_txt = f"⚠ {r[2]}" if es_critico else r[2]
                self.tabla_prod.insert("", "end",
                    values=(
                        r[0], r[1], nombre_txt, f"${r[3]:.2f}", f"${r[4]:.2f}",
                        stock_txt, r[6], r[8] or "-", "Sí" if r[7] else "No", estado
                    ),
                    iid=str(r[0]), tags=tuple(tags))
        self.tabla_prod.tag_configure("low", foreground=C["yellow"])
        self.tabla_prod.tag_configure("critical", foreground=C["red"])

    def _llenar_form_producto(self, event=None):
        sel = self.tabla_prod.selection()
        if not sel:
            return
        vals = self.tabla_prod.item(sel[0], "values")
        pid, codigo, nombre, costo, precio, stock, cat, caducidad, granel, _estado = vals
        nombre = str(nombre).replace("⚠", "").strip()
        costo = costo.replace("$","")
        precio = precio.replace("$","")
        stock = str(stock).replace("kg", "").replace("KG", "").strip()
        caducidad = "" if caducidad == "-" else caducidad
        keys = ("e_codigo","e_nombre","e_costo","e_precio","e_stock","e_categoria","e_caducidad")
        datos = (codigo, nombre, costo, precio, stock, cat, caducidad)
        for k,d in zip(keys, datos):
            self._set_prod_field(k, d)
        self.sv_es_granel.set(str(granel).strip().lower() in ("sí", "si", "1", "true"))
        self._editing_id = int(pid)

    def _guardar_producto(self):
        try:
            codigo   = self._prod_entries["e_codigo"].get().strip()
            nombre   = self._prod_entries["e_nombre"].get().strip()
            costo    = float((self._prod_entries["e_costo"].get().strip() or "0").replace(",", "."))
            precio   = float(self._prod_entries["e_precio"].get().strip().replace(",", "."))
            stock_txt = self._prod_entries["e_stock"].get().strip()
            categoria = self._normalizar_categoria_producto(
                self._prod_entries["e_categoria"].get()
            ) or "General"
            caducidad_txt = self._get_caducidad_form()
        except ValueError:
            messagebox.showerror("Error",
                "Costo y Precio deben ser números válidos.",
                parent=self)
            return
        es_granel = bool(self.sv_es_granel.get())
        try:
            if es_granel:
                stock = float(stock_txt.replace(",", "."))
            else:
                stock = int(stock_txt)
        except ValueError:
            messagebox.showerror(
                "Error",
                "Stock debe ser decimal en kg si es a granel, o entero si no es a granel.",
                parent=self
            )
            return
        caducidad = None
        if caducidad_txt:
            try:
                caducidad = datetime.date.fromisoformat(caducidad_txt).isoformat()
            except ValueError:
                messagebox.showerror(
                    "Error",
                    "La fecha de caducidad debe tener formato YYYY-MM-DD.",
                    parent=self
                )
                return
        if not codigo or not nombre:
            messagebox.showerror("Error", "Código y Nombre son obligatorios.", parent=self)
            return
        if costo < 0 or precio < 0 or stock < 0:
            messagebox.showerror("Error", "Costo, Precio y Stock no pueden ser negativos.", parent=self)
            return
        categoria = self._registrar_categoria_producto(categoria)
        self.sv_categoria_producto.set(categoria)

        eid = getattr(self, "_editing_id", None)
        with get_conn() as conn:
            if eid:
                conn.execute(
                    "UPDATE productos SET codigo=?,nombre=?,costo=?,precio=?,stock=?,categoria=?,a_granel=?,caducidad=?"
                    " WHERE id=?",
                    (codigo, nombre, costo, precio, stock, categoria, int(es_granel), caducidad, eid))
                msg = "Producto actualizado."
            else:
                try:
                    conn.execute(
                        "INSERT INTO productos (codigo,nombre,costo,precio,stock,categoria,a_granel,caducidad)"
                        " VALUES (?,?,?,?,?,?,?,?)",
                        (codigo, nombre, costo, precio, stock, categoria, int(es_granel), caducidad))
                    msg = "Producto agregado."
                except sqlite3.IntegrityError:
                    messagebox.showerror("Error",
                        f'El código "{codigo}" ya existe.', parent=self)
                    return
        messagebox.showinfo("OK", msg, parent=self)
        self._clear_prod_form_inputs()
        self.sv_es_granel.set(False)
        self._editing_id = None
        self._autocodigo()
        self._refrescar_categorias_producto()
        self._cargar_tabla_productos()
        self._cargar_productos()

    def _eliminar_producto(self):
        eid = getattr(self, "_editing_id", None)
        if not eid:
            messagebox.showinfo("Selecciona un producto",
                "Haz clic en un producto de la tabla primero.", parent=self)
            return
        sel = self.tabla_prod.selection()
        nombre = self.tabla_prod.item(sel[0])["values"][2] if sel else "?"
        if messagebox.askyesno("Eliminar",
            f'¿Eliminar "{nombre}"? (No se puede deshacer)', parent=self):
            with get_conn() as conn:
                conn.execute("DELETE FROM productos WHERE id=?", (eid,))
            self._clear_prod_form_inputs()
            self._editing_id = None
            self._autocodigo()
            self._cargar_tabla_productos()
            self._cargar_productos()

    def _autocodigo(self):
        if getattr(self, "_editing_id", None):
            return
        codigo = _generar_codigo_unico()
        e = self._prod_entries["e_codigo"]
        e.delete(0, "end")
        e.insert(0, codigo)

    def _limpiar_form_producto(self):
        self._clear_prod_form_inputs()
        self.sv_es_granel.set(False)
        self._editing_id = None

    def _nuevo_producto(self):
        self._editing_id = None
        self._clear_prod_form_inputs()
        self.sv_es_granel.set(False)
        self._autocodigo()
        if "e_nombre" in self._prod_entries:
            self._prod_entries["e_nombre"].focus()

    def _exportar_inventario_docx(self):
        """Genera el .docx de inventario físico con los datos actuales de la BD
        y lo guarda en la carpeta 'documentos para inventarios fisicos'."""
        # Capa 1: dependencia
        if not _DOCX_OK:
            messagebox.showerror("Dependencia faltante",
                "La librería 'python-docx' no está instalada.\n\n"
                "Instálala con:\n    pip install python-docx", parent=self)
            return
        # Capa 2: leer BD
        try:
            with get_conn() as conn:
                productos = conn.execute(
                    "SELECT id,codigo,nombre,precio,costo,stock,categoria"
                    " FROM productos ORDER BY categoria, nombre"
                ).fetchall()
        except sqlite3.Error as e:
            messagebox.showerror("Error de base de datos",
                f"No se pudo leer la tabla de productos:\n{e}", parent=self)
            return
        if not productos:
            messagebox.showinfo("Sin productos",
                "No hay productos registrados en la base de datos.", parent=self)
            return
        # Capa 3: crear carpeta
        base = os.path.dirname(os.path.abspath(__file__))
        carpeta = os.path.join(base, "documentos para inventarios fisicos")
        try:
            os.makedirs(carpeta, exist_ok=True)
        except OSError as e:
            messagebox.showerror("Error al crear carpeta",
                f"No se pudo crear la carpeta de destino:\n{carpeta}\n\nDetalle: {e}",
                parent=self)
            return
        # Capa 4: generar archivo
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta = os.path.join(carpeta, f"inventario_fisico_{ts}.docx")
        try:
            generar_docx_inventario(productos, ruta)
        except Exception as e:
            messagebox.showerror("Error al generar el documento",
                f"Ocurrió un problema al crear el archivo .docx:\n{e}", parent=self)
            return
        # Capa 5: notificar
        messagebox.showinfo("✔ Inventario generado",
            f"Documento creado correctamente.\n\n"
            f"Productos incluidos: {len(productos)}\n"
            f"Guardado en:\n{ruta}", parent=self)

    # ══════════════════════════════════════════════════════
    #  PÁGINA: CRM
    # ══════════════════════════════════════════════════════
    def _build_page_crm(self):
        page = tk.Frame(self.content, bg=C["bg"])
        self.pages["crm"] = page

        form = tk.Frame(page, bg=C["card"], padx=14, pady=12)
        form.pack(fill="x", pady=(0,10))
        tk.Label(form, text="CLIENTES (CRM LOCAL)", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9, "bold")).grid(row=0, column=0, columnspan=5, sticky="w", pady=(0,8))

        campos = [
            ("Nombre", "c_nombre"),
            ("Teléfono", "c_tel"),
            ("Email", "c_email"),
            ("Notas", "c_notas"),
        ]
        self._cli_entries = {}
        for col, (lbl, key) in enumerate(campos):
            tk.Label(form, text=lbl, fg=C["muted"], bg=C["card"],
                     font=("Courier", 9)).grid(row=1, column=col, sticky="w", padx=(0,4))
            ent = tk.Entry(form, bg=C["panel"], fg=C["text"],
                           insertbackground=C["text"], bd=0,
                           font=("Courier", 11), highlightthickness=1,
                           highlightbackground=C["border"])
            ent.grid(row=2, column=col, sticky="ew", padx=(0,8), ipady=6)
            self._cli_entries[key] = ent

        form.columnconfigure(0, weight=3, minsize=170)
        form.columnconfigure(1, weight=2, minsize=130)
        form.columnconfigure(2, weight=3, minsize=190)
        form.columnconfigure(3, weight=4, minsize=220)

        btns = tk.Frame(form, bg=C["card"])
        btns.grid(row=2, column=4, sticky="e")
        tk.Button(btns, text="● Nuevo", bg=C["accent2"], fg=C["white"], bd=0,
                  font=("Courier", 10), padx=10, pady=6, cursor="hand2",
                  command=self._nuevo_cliente).pack(side="left", padx=(0,4))
        tk.Button(btns, text="＋ Guardar", bg=C["accent"], fg=C["white"], bd=0,
                  font=("Courier", 10, "bold"), padx=10, pady=6, cursor="hand2",
                  command=self._guardar_cliente).pack(side="left", padx=(0,4))
        tk.Button(btns, text="✕ Eliminar", bg=C["red"], fg=C["white"], bd=0,
                  font=("Courier", 10), padx=10, pady=6, cursor="hand2",
                  command=self._eliminar_cliente).pack(side="left")

        filtro = tk.Frame(page, bg=C["bg"])
        filtro.pack(fill="x", pady=(0,6))
        tk.Label(filtro, text="Filtrar cliente:", fg=C["muted"], bg=C["bg"],
                 font=("Courier", 10)).pack(side="left")
        self.sv_cli_filter = tk.StringVar()
        self.sv_cli_filter.trace_add("write", lambda *a: self._cargar_tabla_clientes())
        tk.Entry(filtro, textvariable=self.sv_cli_filter,
                 bg=C["panel"], fg=C["text"], insertbackground=C["text"],
                 bd=0, font=("Courier", 10), highlightthickness=1,
                 highlightbackground=C["border"], width=32).pack(side="left", padx=(8,0), ipady=5)

        frame_t = tk.Frame(page, bg=C["bg"])
        frame_t.pack(fill="both", expand=True)
        cols = ("id", "nombre", "tel", "email", "ultima", "total")
        heads = ("ID", "Nombre", "Teléfono", "Email", "Última visita", "Total compras")
        widths = [45, 200, 130, 200, 130, 120]
        self.tabla_clientes = ttk.Treeview(frame_t, columns=cols, show="headings",
                                           style="POS.Treeview", selectmode="browse")
        for c, h, w in zip(cols, heads, widths):
            self.tabla_clientes.heading(c, text=h, anchor="center")
            self.tabla_clientes.column(c, width=w, anchor="center", stretch=True)
        sb = ttk.Scrollbar(frame_t, orient="vertical", command=self.tabla_clientes.yview)
        self.tabla_clientes.configure(yscrollcommand=sb.set)
        self.tabla_clientes.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.tabla_clientes.bind("<<TreeviewSelect>>", self._llenar_form_cliente)

    def _cargar_tabla_clientes(self):
        if not hasattr(self, "tabla_clientes"):
            return
        q = self.sv_cli_filter.get().strip().lower() if hasattr(self, "sv_cli_filter") else ""
        self._cargar_clientes_cache()
        for row in self.tabla_clientes.get_children():
            self.tabla_clientes.delete(row)
        for r in self._clientes_cache:
            cid, nombre, tel, email, notas, ultima, total = r
            if q and q not in nombre.lower() and q not in (tel or "").lower() and q not in (email or "").lower():
                continue
            ultima_txt = (ultima or "")[:10] if ultima else "-"
            self.tabla_clientes.insert(
                "",
                "end",
                iid=str(cid),
                values=(cid, nombre, tel or "-", email or "-", ultima_txt, f"${(total or 0):.2f}")
            )

    def _llenar_form_cliente(self, event=None):
        if not hasattr(self, "tabla_clientes"):
            return
        sel = self.tabla_clientes.selection()
        if not sel:
            return
        cid = int(sel[0])
        row = next((r for r in self._clientes_cache if r[0] == cid), None)
        if not row:
            return
        _, nombre, tel, email, notas, _ultima, _total = row
        datos = {
            "c_nombre": nombre or "",
            "c_tel": tel or "",
            "c_email": email or "",
            "c_notas": notas or "",
        }
        for key, val in datos.items():
            self._cli_entries[key].delete(0, "end")
            self._cli_entries[key].insert(0, val)
        self._cliente_editing_id = cid

    def _nuevo_cliente(self):
        self._cliente_editing_id = None
        if not hasattr(self, "_cli_entries"):
            return
        for ent in self._cli_entries.values():
            ent.delete(0, "end")
        self._cli_entries["c_nombre"].focus()

    def _guardar_cliente(self):
        if not hasattr(self, "_cli_entries"):
            return
        nombre = self._cli_entries["c_nombre"].get().strip()
        tel = self._cli_entries["c_tel"].get().strip()
        email = self._cli_entries["c_email"].get().strip()
        notas = self._cli_entries["c_notas"].get().strip()
        if not nombre:
            messagebox.showerror("Error", "El nombre del cliente es obligatorio.", parent=self)
            return

        eid = getattr(self, "_cliente_editing_id", None)
        with get_conn() as conn:
            if eid:
                conn.execute(
                    "UPDATE clientes SET nombre=?,telefono=?,email=?,notas=? WHERE id=?",
                    (nombre, tel, email, notas, eid)
                )
                msg = "Cliente actualizado."
            else:
                conn.execute(
                    "INSERT INTO clientes (nombre,telefono,email,notas) VALUES (?,?,?,?)",
                    (nombre, tel, email, notas)
                )
                msg = "Cliente agregado."
        messagebox.showinfo("OK", msg, parent=self)
        self._nuevo_cliente()
        self._cargar_tabla_clientes()
        self._cargar_clientes_en_venta()

    def _eliminar_cliente(self):
        eid = getattr(self, "_cliente_editing_id", None)
        if not eid:
            messagebox.showinfo("Selecciona un cliente", "Primero selecciona un cliente de la tabla.", parent=self)
            return
        nombre = self._cli_entries["c_nombre"].get().strip() or "?"
        if not messagebox.askyesno("Eliminar cliente",
                                   f'¿Eliminar "{nombre}"? Las ventas quedarán como cliente anónimo.',
                                   parent=self):
            return
        with get_conn() as conn:
            conn.execute("UPDATE ventas SET cliente_id = NULL WHERE cliente_id = ?", (eid,))
            conn.execute("DELETE FROM clientes WHERE id = ?", (eid,))
        self._nuevo_cliente()
        self._cargar_tabla_clientes()
        self._cargar_clientes_en_venta()

    # ══════════════════════════════════════════════════════
    #  PÁGINA: CONTABILIDAD
    # ══════════════════════════════════════════════════════
    def _build_page_contabilidad(self):
        page = tk.Frame(self.content, bg=C["bg"])
        self.pages["contabilidad"] = page

        form = tk.Frame(page, bg=C["card"], padx=12, pady=10)
        form.pack(fill="x", pady=(0, 10))
        tk.Label(form, text="MOVIMIENTOS CONTABLES (INGRESOS / EGRESOS)",
                 fg=C["muted"], bg=C["card"], font=("Courier", 9, "bold")).grid(
                     row=0, column=0, columnspan=8, sticky="w", pady=(0, 8))

        self.sv_conta_tipo = tk.StringVar(value="EGRESO")
        self.sv_conta_subtipo = tk.StringVar(value="GASTO")
        self.sv_conta_categoria = tk.StringVar(value="")
        self.sv_conta_metodo = tk.StringVar(value="EFECTIVO")
        self.sv_conta_concepto = tk.StringVar(value="")
        self.sv_conta_monto = tk.StringVar(value="")
        self.sv_conta_ref = tk.StringVar(value="")

        tk.Label(form, text="Tipo", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=0, sticky="w", padx=(0, 4))
        self.cmb_conta_tipo = ttk.Combobox(
            form, textvariable=self.sv_conta_tipo, state="readonly", values=("INGRESO", "EGRESO"), width=10
        )
        self.cmb_conta_tipo.grid(row=2, column=0, sticky="ew", padx=(0, 8), ipady=3)
        self.cmb_conta_tipo.bind("<<ComboboxSelected>>", self._on_conta_tipo_change)

        tk.Label(form, text="Subtipo", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=1, sticky="w", padx=(0, 4))
        self.cmb_conta_subtipo = ttk.Combobox(
            form, textvariable=self.sv_conta_subtipo, state="readonly", width=14
        )
        self.cmb_conta_subtipo.grid(row=2, column=1, sticky="ew", padx=(0, 8), ipady=3)
        self.cmb_conta_subtipo.bind("<<ComboboxSelected>>", self._on_conta_subtipo_change)

        tk.Label(form, text="Categoría", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=2, sticky="w", padx=(0, 4))
        self.cmb_conta_categoria = ttk.Combobox(
            form, textvariable=self.sv_conta_categoria, state="readonly", width=20
        )
        self.cmb_conta_categoria.grid(row=2, column=2, sticky="ew", padx=(0, 8), ipady=3)

        tk.Label(form, text="Concepto", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=3, sticky="w", padx=(0, 4))
        tk.Entry(form, textvariable=self.sv_conta_concepto, bg=C["panel"], fg=C["text"],
                 insertbackground=C["text"], bd=0, font=("Courier", 11),
                 highlightthickness=1, highlightbackground=C["border"]).grid(
                     row=2, column=3, sticky="ew", padx=(0, 8), ipady=6)

        tk.Label(form, text="Monto", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=4, sticky="w", padx=(0, 4))
        tk.Entry(form, textvariable=self.sv_conta_monto, bg=C["panel"], fg=C["text"],
                 insertbackground=C["text"], bd=0, font=("Courier", 11),
                 highlightthickness=1, highlightbackground=C["border"], width=12).grid(
                     row=2, column=4, sticky="ew", padx=(0, 8), ipady=6)

        tk.Label(form, text="Método", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=5, sticky="w", padx=(0, 4))
        self.cmb_conta_metodo = ttk.Combobox(
            form, textvariable=self.sv_conta_metodo, state="readonly",
            values=("EFECTIVO", "TRANSFERENCIA", "TARJETA", "OTRO"), width=14
        )
        self.cmb_conta_metodo.grid(row=2, column=5, sticky="ew", padx=(0, 8), ipady=3)

        tk.Label(form, text="Referencia", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=6, sticky="w", padx=(0, 4))
        tk.Entry(form, textvariable=self.sv_conta_ref, bg=C["panel"], fg=C["text"],
                 insertbackground=C["text"], bd=0, font=("Courier", 11),
                 highlightthickness=1, highlightbackground=C["border"], width=14).grid(
                     row=2, column=6, sticky="ew", padx=(0, 8), ipady=6)

        btns = tk.Frame(form, bg=C["card"])
        btns.grid(row=2, column=7, sticky="e")
        tk.Button(btns, text="＋ Registrar", bg=C["accent"], fg=C["white"], bd=0,
                  font=("Courier", 10, "bold"), padx=10, pady=6, cursor="hand2",
                  command=self._registrar_movimiento_desde_form).pack(side="left", padx=(0, 4))
        tk.Button(btns, text="⟳ Limpiar", bg=C["panel"], fg=C["muted"], bd=0,
                  font=("Courier", 10), padx=10, pady=6, cursor="hand2",
                  command=self._limpiar_form_contabilidad).pack(side="left")

        form.columnconfigure(0, weight=1, minsize=90)
        form.columnconfigure(1, weight=1, minsize=110)
        form.columnconfigure(2, weight=2, minsize=140)
        form.columnconfigure(3, weight=4, minsize=240)
        form.columnconfigure(4, weight=1, minsize=100)
        form.columnconfigure(5, weight=1, minsize=120)
        form.columnconfigure(6, weight=1, minsize=120)

        filt = tk.Frame(page, bg=C["card"], padx=12, pady=10)
        filt.pack(fill="x", pady=(0, 10))
        tk.Label(filt, text="Desde (YYYY-MM-DD):", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).pack(side="left")
        hoy = datetime.date.today().isoformat()
        self.sv_conta_ini = tk.StringVar(value=hoy)
        self.sv_conta_fin = tk.StringVar(value=hoy)
        tk.Entry(filt, textvariable=self.sv_conta_ini, bg=C["panel"], fg=C["text"],
                 insertbackground=C["text"], bd=0, font=("Courier", 10),
                 highlightthickness=1, highlightbackground=C["border"], width=12).pack(
                     side="left", padx=(6, 10), ipady=5)
        tk.Label(filt, text="Hasta:", fg=C["muted"], bg=C["card"],
                 font=("Courier", 9)).pack(side="left")
        tk.Entry(filt, textvariable=self.sv_conta_fin, bg=C["panel"], fg=C["text"],
                 insertbackground=C["text"], bd=0, font=("Courier", 10),
                 highlightthickness=1, highlightbackground=C["border"], width=12).pack(
                     side="left", padx=(6, 10), ipady=5)
        tk.Button(filt, text="Hoy", bg=C["panel"], fg=C["muted"], bd=0,
                  font=("Courier", 10), padx=10, pady=5, cursor="hand2",
                  command=self._conta_set_hoy).pack(side="left")
        tk.Button(filt, text="Actualizar", bg=C["accent2"], fg=C["white"], bd=0,
                  font=("Courier", 10), padx=10, pady=5, cursor="hand2",
                  command=self._cargar_tabla_contabilidad).pack(side="right", padx=(8, 0))
        tk.Button(filt, text="💾 Cortes CSV", bg=C["accent2"], fg=C["white"], bd=0,
                  font=("Courier", 10), padx=10, pady=5, cursor="hand2",
                  command=self._exportar_cortes_caja_csv_ui).pack(side="right", padx=(8, 0))
        tk.Button(filt, text="💾 Resumen mensual CSV", bg=C["accent2"], fg=C["white"], bd=0,
                  font=("Courier", 10), padx=10, pady=5, cursor="hand2",
                  command=self._exportar_resumen_mensual_csv_ui).pack(side="right", padx=(8, 0))
        tk.Button(filt, text="💾 Libro Diario CSV", bg=C["accent"], fg=C["white"], bd=0,
                  font=("Courier", 10, "bold"), padx=10, pady=5, cursor="hand2",
                  command=self._exportar_libro_diario_csv_ui).pack(side="right")

        caja = tk.Frame(page, bg=C["card"], padx=12, pady=10)
        caja.pack(fill="x", pady=(0, 10))
        tk.Label(caja, text="APERTURA / CIERRE DE CAJA",
                 fg=C["muted"], bg=C["card"], font=("Courier", 9, "bold")).grid(
                     row=0, column=0, columnspan=8, sticky="w", pady=(0, 8))

        self.sv_caja_estado = tk.StringVar(value="CERRADO")
        self.sv_caja_id = tk.StringVar(value="-")
        self.sv_caja_apertura = tk.StringVar(value="-")
        self.sv_caja_ventas = tk.StringVar(value="$0.00")
        self.sv_caja_ingresos_extra = tk.StringVar(value="$0.00")
        self.sv_caja_egresos = tk.StringVar(value="$0.00")
        self.sv_caja_saldo_teorico = tk.StringVar(value="$0.00")
        self.sv_caja_saldo_inicial = tk.StringVar(value="0.00")
        self.sv_caja_saldo_real = tk.StringVar(value="")
        self.sv_caja_diferencia = tk.StringVar(value="$0.00")
        self._caja_abierta_id = None
        self._caja_saldo_teorico_actual = 0.0
        self.sv_caja_saldo_real.trace_add("write", lambda *a: self._conta_recalcular_diferencia_preview())

        tk.Label(caja, text="Estado", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=0, sticky="w", padx=(0, 4))
        self.lbl_caja_estado = tk.Label(
            caja, textvariable=self.sv_caja_estado, fg=C["red"], bg=C["card"], font=("Courier", 11, "bold")
        )
        self.lbl_caja_estado.grid(row=2, column=0, sticky="w", padx=(0, 10))

        tk.Label(caja, text="Folio", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=1, sticky="w", padx=(0, 4))
        tk.Label(caja, textvariable=self.sv_caja_id, fg=C["text"], bg=C["card"], font=("Courier", 11)).grid(
            row=2, column=1, sticky="w", padx=(0, 10))

        tk.Label(caja, text="Apertura", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=2, sticky="w", padx=(0, 4))
        tk.Label(caja, textvariable=self.sv_caja_apertura, fg=C["text"], bg=C["card"], font=("Courier", 10)).grid(
            row=2, column=2, sticky="w", padx=(0, 10))

        tk.Label(caja, text="Ventas efectivo", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=3, sticky="w", padx=(0, 4))
        tk.Label(caja, textvariable=self.sv_caja_ventas, fg=C["green"], bg=C["card"], font=("Courier", 11, "bold")).grid(
            row=2, column=3, sticky="w", padx=(0, 10))

        tk.Label(caja, text="Ingresos extra", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=4, sticky="w", padx=(0, 4))
        tk.Label(caja, textvariable=self.sv_caja_ingresos_extra, fg=C["green"], bg=C["card"], font=("Courier", 11, "bold")).grid(
            row=2, column=4, sticky="w", padx=(0, 10))

        tk.Label(caja, text="Egresos", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=5, sticky="w", padx=(0, 4))
        tk.Label(caja, textvariable=self.sv_caja_egresos, fg=C["red"], bg=C["card"], font=("Courier", 11, "bold")).grid(
            row=2, column=5, sticky="w", padx=(0, 10))

        tk.Label(caja, text="Saldo teórico", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=1, column=6, sticky="w", padx=(0, 4))
        tk.Label(caja, textvariable=self.sv_caja_saldo_teorico, fg=C["accent"], bg=C["card"], font=("Courier", 11, "bold")).grid(
            row=2, column=6, sticky="w", padx=(0, 10))

        tk.Label(caja, text="Saldo inicial", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=3, column=0, sticky="w", padx=(0, 4), pady=(8, 0))
        tk.Entry(caja, textvariable=self.sv_caja_saldo_inicial, bg=C["panel"], fg=C["text"],
                 insertbackground=C["text"], bd=0, font=("Courier", 11),
                 highlightthickness=1, highlightbackground=C["border"], width=12).grid(
                     row=4, column=0, sticky="ew", padx=(0, 8), ipady=6)

        tk.Label(caja, text="Saldo real (cierre)", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=3, column=1, sticky="w", padx=(0, 4), pady=(8, 0))
        tk.Entry(caja, textvariable=self.sv_caja_saldo_real, bg=C["panel"], fg=C["text"],
                 insertbackground=C["text"], bd=0, font=("Courier", 11),
                 highlightthickness=1, highlightbackground=C["border"], width=12).grid(
                     row=4, column=1, sticky="ew", padx=(0, 8), ipady=6)

        tk.Label(caja, text="Diferencia", fg=C["muted"], bg=C["card"], font=("Courier", 9)).grid(
            row=3, column=2, sticky="w", padx=(0, 4), pady=(8, 0))
        tk.Entry(caja, textvariable=self.sv_caja_diferencia, state="readonly", readonlybackground=C["panel"],
                 fg=C["text"], bd=0, font=("Courier", 11), width=14).grid(
                     row=4, column=2, sticky="ew", padx=(0, 8), ipady=6)

        btns_caja = tk.Frame(caja, bg=C["card"])
        btns_caja.grid(row=4, column=3, columnspan=5, sticky="e")
        self.btn_caja_abrir = tk.Button(
            btns_caja, text="Abrir caja", bg=C["accent"], fg=C["white"], bd=0,
            font=("Courier", 10, "bold"), padx=10, pady=6, cursor="hand2",
            command=self._conta_abrir_caja_ui
        )
        self.btn_caja_abrir.pack(side="left", padx=(0, 6))
        tk.Button(btns_caja, text="Actualizar resumen", bg=C["panel"], fg=C["muted"], bd=0,
                  font=("Courier", 10), padx=10, pady=6, cursor="hand2",
                  command=self._conta_refrescar_resumen_caja_ui).pack(side="left", padx=(0, 6))
        self.btn_caja_cerrar = tk.Button(
            btns_caja, text="Cerrar caja", bg=C["red"], fg=C["white"], bd=0,
            font=("Courier", 10, "bold"), padx=10, pady=6, cursor="hand2",
            command=self._conta_cerrar_caja_ui
        )
        self.btn_caja_cerrar.pack(side="left")

        for idx in range(8):
            caja.columnconfigure(idx, weight=1)
        caja.columnconfigure(2, weight=2)
        caja.columnconfigure(7, weight=1)

        kpi = tk.Frame(page, bg=C["bg"])
        kpi.pack(fill="x", pady=(0, 8))
        self.kpi_ingresos = self._kpi_box(kpi, "INGRESOS", "$0.00", C["green"])
        self.kpi_egresos = self._kpi_box(kpi, "EGRESOS", "$0.00", C["red"])
        self.kpi_flujo = self._kpi_box(kpi, "FLUJO NETO", "$0.00", C["accent2"])

        frame_t = tk.Frame(page, bg=C["bg"])
        frame_t.pack(fill="both", expand=True)
        cols = ("id", "fecha", "tipo", "subtipo", "categoria", "concepto", "metodo", "monto", "venta", "ref")
        heads = ("Folio", "Fecha", "Tipo", "Subtipo", "Categoría", "Concepto", "Método", "Monto", "Venta", "Referencia")
        widths = [55, 145, 85, 115, 140, 230, 105, 95, 65, 120]
        self.tabla_conta = ttk.Treeview(frame_t, columns=cols, show="headings",
                                        style="POS.Treeview", selectmode="browse")
        for c, h, w in zip(cols, heads, widths):
            self.tabla_conta.heading(c, text=h, anchor="center")
            self.tabla_conta.column(c, width=w, anchor="center", stretch=True)
        sb = ttk.Scrollbar(frame_t, orient="vertical", command=self.tabla_conta.yview)
        self.tabla_conta.configure(yscrollcommand=sb.set)
        self.tabla_conta.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.tabla_conta.tag_configure("ing", foreground=C["green"])
        self.tabla_conta.tag_configure("egr", foreground=C["red"])

        self._cargar_contabilidad_vista()

    def _conta_set_hoy(self):
        hoy = datetime.date.today().isoformat()
        if hasattr(self, "sv_conta_ini"):
            self.sv_conta_ini.set(hoy)
        if hasattr(self, "sv_conta_fin"):
            self.sv_conta_fin.set(hoy)
        self._cargar_tabla_contabilidad()

    def _conta_recalcular_diferencia_preview(self):
        if not hasattr(self, "sv_caja_diferencia"):
            return
        txt_real = self.sv_caja_saldo_real.get().strip()
        if not txt_real:
            self.sv_caja_diferencia.set("$0.00")
            return
        try:
            saldo_real = float(txt_real.replace(",", "."))
        except ValueError:
            self.sv_caja_diferencia.set("Inválido")
            return
        diferencia = saldo_real - float(getattr(self, "_caja_saldo_teorico_actual", 0.0) or 0.0)
        self.sv_caja_diferencia.set(f"${diferencia:.2f}")

    def _conta_refrescar_resumen_caja_ui(self):
        if not hasattr(self, "sv_caja_estado"):
            return

        resumen = self.resumen_caja_actual()
        if resumen:
            self._caja_abierta_id = resumen["id"]
            self.sv_caja_estado.set("ABIERTO")
            if hasattr(self, "lbl_caja_estado"):
                self.lbl_caja_estado.config(fg=C["green"])
            self.sv_caja_id.set(f"#{resumen['id']}")
            self.sv_caja_apertura.set(resumen["fecha_apertura"])
            self.sv_caja_ventas.set(f"${float(resumen['ventas_efectivo']):.2f}")
            self.sv_caja_ingresos_extra.set(f"${float(resumen['ingresos_extra']):.2f}")
            self.sv_caja_egresos.set(f"${float(resumen['egresos']):.2f}")
            self.sv_caja_saldo_teorico.set(f"${float(resumen['saldo_teorico']):.2f}")
            self._caja_saldo_teorico_actual = float(resumen["saldo_teorico"])
            if hasattr(self, "btn_caja_abrir"):
                self.btn_caja_abrir.config(state="disabled")
            if hasattr(self, "btn_caja_cerrar"):
                self.btn_caja_cerrar.config(state="normal")
            self._conta_recalcular_diferencia_preview()
            return

        with get_conn() as conn:
            ultimo = conn.execute(
                "SELECT id,fecha_apertura,saldo_inicial,ventas_efectivo,ingresos_extra,"
                " egresos,saldo_teorico,saldo_real,diferencia,estado"
                " FROM cortes_caja ORDER BY id DESC LIMIT 1"
            ).fetchone()

        self._caja_abierta_id = None
        if hasattr(self, "btn_caja_abrir"):
            self.btn_caja_abrir.config(state="normal")
        if hasattr(self, "btn_caja_cerrar"):
            self.btn_caja_cerrar.config(state="disabled")
        if hasattr(self, "lbl_caja_estado"):
            self.lbl_caja_estado.config(fg=C["red"])

        if not ultimo:
            self.sv_caja_estado.set("CERRADO")
            self.sv_caja_id.set("-")
            self.sv_caja_apertura.set("-")
            self.sv_caja_ventas.set("$0.00")
            self.sv_caja_ingresos_extra.set("$0.00")
            self.sv_caja_egresos.set("$0.00")
            self.sv_caja_saldo_teorico.set("$0.00")
            self.sv_caja_diferencia.set("$0.00")
            self._caja_saldo_teorico_actual = 0.0
            return

        self.sv_caja_estado.set(str(ultimo[9] or "CERRADO"))
        self.sv_caja_id.set(f"#{ultimo[0]}")
        self.sv_caja_apertura.set(ultimo[1] or "-")
        self.sv_caja_saldo_inicial.set(f"{float(ultimo[2] or 0):.2f}")
        self.sv_caja_ventas.set(f"${float(ultimo[3] or 0):.2f}")
        self.sv_caja_ingresos_extra.set(f"${float(ultimo[4] or 0):.2f}")
        self.sv_caja_egresos.set(f"${float(ultimo[5] or 0):.2f}")
        self.sv_caja_saldo_teorico.set(f"${float(ultimo[6] or 0):.2f}")
        self._caja_saldo_teorico_actual = float(ultimo[6] or 0)
        if ultimo[7] is not None:
            self.sv_caja_saldo_real.set(f"{float(ultimo[7]):.2f}")
        else:
            self.sv_caja_saldo_real.set("")
        if ultimo[8] is not None:
            self.sv_caja_diferencia.set(f"${float(ultimo[8]):.2f}")
        else:
            self._conta_recalcular_diferencia_preview()

    def _conta_abrir_caja_ui(self):
        try:
            saldo_inicial = float((self.sv_caja_saldo_inicial.get().strip() or "0").replace(",", "."))
        except ValueError:
            messagebox.showerror("Error", "El saldo inicial debe ser un número válido.", parent=self)
            return

        try:
            corte_id = self.abrir_caja(saldo_inicial, usuario="admin")
        except (ValueError, RuntimeError) as e:
            messagebox.showerror("Caja", str(e), parent=self)
            return
        except sqlite3.Error as e:
            messagebox.showerror("Error de base de datos",
                                 f"No se pudo abrir la caja.\nDetalle: {e}", parent=self)
            return

        self.sv_caja_saldo_real.set("")
        self.sv_caja_diferencia.set("$0.00")
        self._conta_refrescar_resumen_caja_ui()
        messagebox.showinfo("Caja abierta",
                            f"Caja #{corte_id} abierta con saldo inicial ${saldo_inicial:.2f}.",
                            parent=self)

    def _conta_cerrar_caja_ui(self):
        try:
            saldo_real = float((self.sv_caja_saldo_real.get().strip() or "0").replace(",", "."))
        except ValueError:
            messagebox.showerror("Error", "El saldo real debe ser un número válido.", parent=self)
            return
        if saldo_real < 0:
            messagebox.showerror("Error", "El saldo real no puede ser negativo.", parent=self)
            return

        saldo_teorico = float(getattr(self, "_caja_saldo_teorico_actual", 0.0) or 0.0)
        try:
            diferencia = self.cerrar_caja(saldo_real, usuario="admin")
        except (ValueError, RuntimeError) as e:
            messagebox.showerror("Caja", str(e), parent=self)
            return
        except sqlite3.Error as e:
            messagebox.showerror("Error de base de datos",
                                 f"No se pudo cerrar la caja.\nDetalle: {e}", parent=self)
            return

        self._conta_refrescar_resumen_caja_ui()
        messagebox.showinfo(
            "Caja cerrada",
            f"Saldo teórico: ${saldo_teorico:.2f}\n"
            f"Saldo real: ${saldo_real:.2f}\n"
            f"Diferencia: ${float(diferencia):.2f}",
            parent=self
        )

    def _rango_fechas_conta(self):
        ini = datetime.date.fromisoformat(self.sv_conta_ini.get().strip())
        fin = datetime.date.fromisoformat(self.sv_conta_fin.get().strip())
        if fin < ini:
            raise ValueError("La fecha final no puede ser menor a la inicial.")
        return ini.isoformat(), fin.isoformat(), f"{ini} 00:00:00", f"{fin} 23:59:59"

    def _on_conta_tipo_change(self, event=None):
        tipo = self.sv_conta_tipo.get().strip().upper()
        subtipos = self._opciones_subtipo_por_tipo(tipo)
        self.cmb_conta_subtipo["values"] = tuple(subtipos)
        if self.sv_conta_subtipo.get() not in subtipos:
            self.sv_conta_subtipo.set(subtipos[0] if subtipos else "")
        self._on_conta_subtipo_change()

    def _on_conta_subtipo_change(self, event=None):
        tipo = self.sv_conta_tipo.get().strip().upper()
        subtipo = self.sv_conta_subtipo.get().strip().upper()
        cats = self._categorias_por_tipo(tipo)
        self.cmb_conta_categoria["values"] = tuple(cats)
        preferida = self._categoria_default(tipo, subtipo)
        if preferida in cats:
            self.sv_conta_categoria.set(preferida)
        elif cats and self.sv_conta_categoria.get() not in cats:
            self.sv_conta_categoria.set(cats[0])
        elif not cats:
            self.sv_conta_categoria.set("")

    def _limpiar_form_contabilidad(self):
        self.sv_conta_tipo.set("EGRESO")
        self.sv_conta_subtipo.set("GASTO")
        self.sv_conta_categoria.set("")
        self.sv_conta_metodo.set("EFECTIVO")
        self.sv_conta_concepto.set("")
        self.sv_conta_monto.set("")
        self.sv_conta_ref.set("")
        self._on_conta_tipo_change()

    def _registrar_movimiento_desde_form(self):
        try:
            tipo = self.sv_conta_tipo.get().strip().upper()
            subtipo = self.sv_conta_subtipo.get().strip().upper()
            categoria = self.sv_conta_categoria.get().strip()
            concepto = self.sv_conta_concepto.get().strip()
            metodo = self.sv_conta_metodo.get().strip().upper() or "EFECTIVO"
            referencia = self.sv_conta_ref.get().strip()
            monto = float((self.sv_conta_monto.get().strip() or "0").replace(",", "."))
        except ValueError:
            messagebox.showerror("Error", "El monto debe ser un número válido.", parent=self)
            return

        try:
            if subtipo == "GASTO" and tipo == "EGRESO":
                self.registrar_gasto(categoria, concepto, monto, referencia=referencia)
            elif subtipo == "RETIRO_DUENO" and tipo == "EGRESO":
                self.registrar_retiro_dueno(monto, concepto or "Retiro de dueño")
            elif subtipo == "INGRESO_EXTRA" and tipo == "INGRESO":
                self.registrar_ingreso_extra(categoria, concepto, monto)
            else:
                self._registrar_movimiento_caja(
                    tipo=tipo, subtipo=subtipo, categoria=categoria, concepto=concepto,
                    monto=monto, metodo_pago=metodo, referencia=referencia
                )
        except ValueError as e:
            messagebox.showerror("Error", str(e), parent=self)
            return
        except sqlite3.Error as e:
            messagebox.showerror("Error de base de datos",
                                 f"No se pudo registrar el movimiento.\nDetalle: {e}", parent=self)
            return

        messagebox.showinfo("OK", "Movimiento contable registrado.", parent=self)
        self.sv_conta_concepto.set("")
        self.sv_conta_monto.set("")
        self.sv_conta_ref.set("")
        self._cargar_tabla_contabilidad()

    def _cargar_contabilidad_vista(self):
        if not hasattr(self, "tabla_conta"):
            return
        self._cargar_categorias_contables_cache()
        self._on_conta_tipo_change()
        self._cargar_tabla_contabilidad()

    def _cargar_tabla_contabilidad(self):
        if not hasattr(self, "tabla_conta"):
            return
        try:
            _ini, _fin, inicio, fin = self._rango_fechas_conta()
        except ValueError as e:
            messagebox.showerror("Fechas inválidas", str(e), parent=self)
            return

        for row in self.tabla_conta.get_children():
            self.tabla_conta.delete(row)

        with get_conn() as conn:
            rows = conn.execute(
                "SELECT id,fecha,tipo,subtipo,categoria,concepto,metodo_pago,monto,venta_id,referencia"
                " FROM movimientos_caja"
                " WHERE fecha >= ? AND fecha <= ?"
                " ORDER BY id DESC LIMIT 500",
                (inicio, fin)
            ).fetchall()
            sums = conn.execute(
                "SELECT"
                " IFNULL(SUM(CASE WHEN tipo='INGRESO' THEN monto END),0),"
                " IFNULL(SUM(CASE WHEN tipo='EGRESO' THEN monto END),0)"
                " FROM movimientos_caja WHERE fecha >= ? AND fecha <= ?",
                (inicio, fin)
            ).fetchone()

        for r in rows:
            tag = "ing" if r[2] == "INGRESO" else "egr"
            monto_txt = f"${float(r[7]):.2f}"
            venta_txt = r[8] if r[8] is not None else "-"
            self.tabla_conta.insert(
                "", "end", iid=str(r[0]),
                values=(r[0], r[1], r[2], r[3], r[4], r[5], r[6], monto_txt, venta_txt, r[9] or "-"),
                tags=(tag,)
            )

        ingresos = float(sums[0] or 0)
        egresos = float(sums[1] or 0)
        flujo = ingresos - egresos
        self.kpi_ingresos.config(text=f"${ingresos:.2f}")
        self.kpi_egresos.config(text=f"${egresos:.2f}")
        self.kpi_flujo.config(text=f"${flujo:.2f}")
        self._conta_refrescar_resumen_caja_ui()

    def _exportar_libro_diario_csv_ui(self):
        try:
            fecha_ini, fecha_fin, _inicio, _fin = self._rango_fechas_conta()
        except ValueError as e:
            messagebox.showerror("Fechas inválidas", str(e), parent=self)
            return

        base = os.path.dirname(os.path.abspath(__file__))
        carpeta = os.path.join(base, "respaldos_csv", "contabilidad")
        os.makedirs(carpeta, exist_ok=True)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta = os.path.join(carpeta, f"libro_diario_{fecha_ini}_{fecha_fin}_{ts}.csv")
        try:
            total = self.exportar_libro_diario_csv(fecha_ini, fecha_fin, ruta)
        except Exception as e:
            messagebox.showerror("Error de exportación",
                                 f"No se pudo exportar el libro diario.\nDetalle: {e}", parent=self)
            return
        messagebox.showinfo(
            "✔ Exportación contable",
            f"Movimientos exportados: {total}\n\nArchivo:\n{ruta}",
            parent=self
        )

    def _exportar_resumen_mensual_csv_ui(self):
        try:
            _fecha_ini, fecha_fin, _inicio, _fin = self._rango_fechas_conta()
            fecha_ref = datetime.date.fromisoformat(fecha_fin)
        except ValueError as e:
            messagebox.showerror("Fechas inválidas", str(e), parent=self)
            return

        base = os.path.dirname(os.path.abspath(__file__))
        carpeta = os.path.join(base, "respaldos_csv", "contabilidad")
        os.makedirs(carpeta, exist_ok=True)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        anio, mes = fecha_ref.year, fecha_ref.month
        ruta = os.path.join(carpeta, f"resumen_mensual_{anio:04d}-{mes:02d}_{ts}.csv")

        try:
            total = self.exportar_resumen_mensual_csv(anio, mes, ruta)
        except Exception as e:
            messagebox.showerror("Error de exportación",
                                 f"No se pudo exportar el resumen mensual.\nDetalle: {e}",
                                 parent=self)
            return

        messagebox.showinfo(
            "✔ Exportación contable",
            f"Mes exportado: {anio:04d}-{mes:02d}\n"
            f"Registros: {total}\n\n"
            f"Archivo:\n{ruta}",
            parent=self
        )

    def _exportar_cortes_caja_csv_ui(self):
        try:
            fecha_ini, fecha_fin, _inicio, _fin = self._rango_fechas_conta()
        except ValueError as e:
            messagebox.showerror("Fechas inválidas", str(e), parent=self)
            return

        base = os.path.dirname(os.path.abspath(__file__))
        carpeta = os.path.join(base, "respaldos_csv", "contabilidad")
        os.makedirs(carpeta, exist_ok=True)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        ruta = os.path.join(carpeta, f"cortes_caja_{fecha_ini}_{fecha_fin}_{ts}.csv")
        try:
            total = self.exportar_cortes_caja_csv(fecha_ini, fecha_fin, ruta)
        except Exception as e:
            messagebox.showerror("Error de exportación",
                                 f"No se pudo exportar cortes de caja.\nDetalle: {e}", parent=self)
            return
        messagebox.showinfo(
            "✔ Exportación contable",
            f"Cortes exportados: {total}\n\nArchivo:\n{ruta}",
            parent=self
        )

    # ══════════════════════════════════════════════════════
    #  PÁGINA: HISTORIAL
    # ══════════════════════════════════════════════════════
    def _build_page_historial(self):
        page = tk.Frame(self.content, bg=C["bg"])
        self.pages["historial"] = page

        filter_f = tk.Frame(page, bg=C["card"], padx=12, pady=10)
        filter_f.pack(fill="x", pady=(0,10))
        tk.Label(filter_f, text="Filtrar por fecha (YYYY-MM-DD):", fg=C["muted"],
                 bg=C["card"], font=("Courier",9)).pack(side="left")
        self.sv_hist_fecha = tk.StringVar()
        self.sv_hist_fecha.trace_add("write", lambda *a: self._cargar_historial())
        tk.Entry(filter_f, textvariable=self.sv_hist_fecha,
                 bg=C["panel"], fg=C["text"], insertbackground=C["text"],
                 bd=0, font=("Courier",11), width=14, highlightthickness=1,
                 highlightbackground=C["border"]).pack(side="left", padx=8, ipady=5)

        tk.Button(filter_f, text="🗑  Eliminar venta seleccionada",
                  bg=C["red"], fg=C["white"], bd=0,
                  font=("Courier", 10, "bold"), padx=14, pady=4, cursor="hand2",
                  activebackground="#c0392b",
                  command=self._eliminar_venta_protegida).pack(side="right")

        tk.Button(filter_f, text="🔑  Cambiar contraseña",
                  bg=C["accent2"], fg=C["white"], bd=0,
                  font=("Courier", 10), padx=12, pady=4, cursor="hand2",
                  activebackground="#6a4aaf",
                  command=self._cambiar_contrasena).pack(side="right", padx=(0,8))

        tk.Button(filter_f, text="💾  Exportar CSV",
                  bg=C["accent"], fg=C["white"], bd=0,
                  font=("Courier", 10, "bold"), padx=12, pady=4, cursor="hand2",
                  activebackground="#3a7de0",
                  command=self._exportar_csv).pack(side="right", padx=(0,8))

        self.kpi_frame = tk.Frame(page, bg=C["bg"])
        self.kpi_frame.pack(fill="x", pady=(0,10))
        self.kpi_ventas   = self._kpi_box(self.kpi_frame, "VENTAS HOY",   "0",      C["accent"])
        self.kpi_total    = self._kpi_box(self.kpi_frame, "TOTAL HOY",    "$0.00",  C["green"])
        self.kpi_ganancia = self._kpi_box(self.kpi_frame, "GANANCIA HOY", "$0.00",  C["yellow"])

        cols_v = ("id","fecha","cliente","total")
        fr1 = tk.Frame(page, bg=C["bg"])
        fr1.pack(fill="both", expand=True)

        tk.Label(fr1, text="VENTAS REGISTRADAS", fg=C["muted"], bg=C["bg"],
                 font=("Courier",9,"bold")).pack(anchor="w", pady=(0,4))

        split = tk.Frame(fr1, bg=C["bg"])
        split.pack(fill="both", expand=True)

        left_h = tk.Frame(split, bg=C["bg"])
        left_h.pack(side="left", fill="both", expand=True, padx=(0,8))

        self.tabla_hist = ttk.Treeview(left_h, columns=cols_v, show="headings",
                                       style="POS.Treeview", height=8)
        for c,h,w in zip(cols_v,("ID","Fecha","Cliente","Total"),(40,150,170,80)):
            self.tabla_hist.heading(c, text=h, anchor="center")
            self.tabla_hist.column(c, width=w, anchor="center", stretch=True)
        sb = ttk.Scrollbar(left_h, orient="vertical", command=self.tabla_hist.yview)
        self.tabla_hist.configure(yscrollcommand=sb.set)
        self.tabla_hist.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.tabla_hist.bind("<<TreeviewSelect>>", self._ver_detalle_venta)

        right_h = tk.Frame(split, bg=C["card"], padx=12, pady=12, width=480)
        right_h.pack(side="right", fill="y")
        right_h.pack_propagate(False)

        tk.Label(right_h, text="DETALLE DE VENTA", fg=C["muted"], bg=C["card"],
                 font=("Courier",9,"bold")).pack(anchor="w", pady=(0,6))

        self.tabla_det = ttk.Treeview(right_h,
            columns=("nombre","cant","precio","sub","ganancia"),
            show="headings", style="POS.Treeview")
        for c,h,w in zip(("nombre","cant","precio","sub","ganancia"),
                          ("Producto","Cant","Precio","Subtotal","Ganancia"),
                          (140,50,85,90,90)):
            self.tabla_det.heading(c, text=h, anchor="center")
            self.tabla_det.column(c, width=w, anchor="center", stretch=False)
        self.tabla_det.pack(fill="both", expand=True)

        def _block_resize(event, tv):
            if tv.identify_region(event.x, event.y) == "separator":
                return "break"
        self.tabla_hist.bind("<Button-1>", lambda e: _block_resize(e, self.tabla_hist), add=True)
        self.tabla_hist.bind("<B1-Motion>", lambda e: _block_resize(e, self.tabla_hist), add=True)
        self.tabla_det.bind("<Button-1>", lambda e: _block_resize(e, self.tabla_det), add=True)
        self.tabla_det.bind("<B1-Motion>", lambda e: _block_resize(e, self.tabla_det), add=True)

    def _kpi_box(self, parent, label, value, color):
        box = tk.Frame(parent, bg=C["card"], padx=20, pady=12)
        box.pack(side="left", padx=(0,8))
        tk.Label(box, text=label, fg=C["muted"], bg=C["card"],
                 font=("Courier",8)).pack(anchor="w")
        lbl = tk.Label(box, text=value, fg=color, bg=C["card"],
                       font=("Courier",18,"bold"))
        lbl.pack(anchor="w")
        return lbl

    def _cargar_historial(self):
        f = ""
        if hasattr(self, "sv_hist_fecha"):
            f = self.sv_hist_fecha.get().strip()
        today = datetime.date.today().isoformat()
        for row in self.tabla_hist.get_children():
            self.tabla_hist.delete(row)
        with get_conn() as conn:
            if f:
                rows = conn.execute(
                    "SELECT v.id,v.fecha,IFNULL(c.nombre,'Público general') AS cliente,v.total"
                    " FROM ventas v"
                    " LEFT JOIN clientes c ON c.id = v.cliente_id"
                    " WHERE v.fecha LIKE ?"
                    " ORDER BY v.id DESC",
                    (f"%{f}%",)).fetchall()
            else:
                rows = conn.execute(
                    "SELECT v.id,v.fecha,IFNULL(c.nombre,'Público general') AS cliente,v.total"
                    " FROM ventas v"
                    " LEFT JOIN clientes c ON c.id = v.cliente_id"
                    " ORDER BY v.id DESC LIMIT 200"
                ).fetchall()
            kpi = conn.execute(
                "SELECT COUNT(*), IFNULL(SUM(v.total),0),"
                " IFNULL(SUM(dv.ganancia),0)"
                " FROM ventas v"
                " LEFT JOIN detalle_venta dv ON dv.venta_id = v.id"
                " WHERE v.fecha LIKE ?",
                (f"{today}%",)).fetchone()
        for r in rows:
            self.tabla_hist.insert("","end",
                values=(r[0],r[1],r[2],f"${r[3]:.2f}"), iid=str(r[0]))
        if hasattr(self,"kpi_ventas"):
            self.kpi_ventas.config(text=str(kpi[0]))
            self.kpi_total.config(text=f"${kpi[1]:.2f}")
            self.kpi_ganancia.config(text=f"${kpi[2]:.2f}")

    def _ver_detalle_venta(self, event=None):
        sel = self.tabla_hist.selection()
        if not sel:
            return
        vid = int(sel[0])
        for row in self.tabla_det.get_children():
            self.tabla_det.delete(row)
        with get_conn() as conn:
            rows = conn.execute(
                "SELECT nombre,cantidad,precio,subtotal,ganancia,es_granel FROM detalle_venta"
                " WHERE venta_id=?", (vid,)).fetchall()
        for r in rows:
            cant_txt = self._fmt_unidades(r[1], bool(r[5]), bool(r[5]))
            self.tabla_det.insert("","end",
                values=(r[0],cant_txt,f"${r[2]:.2f}",f"${r[3]:.2f}",f"${r[4]:.2f}"))

    def _exportar_csv(self):
        carpeta_base = os.path.dirname(os.path.abspath(__file__))
        carpeta_respaldos = os.path.join(carpeta_base, "respaldos_csv")
        os.makedirs(carpeta_respaldos, exist_ok=True)
        fecha_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        nombre_archivo = f"historial_ventas_{fecha_str}.csv"
        ruta_completa = os.path.join(carpeta_respaldos, nombre_archivo)
        try:
            with get_conn() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT v.id, v.fecha, IFNULL(c.nombre,'Público general') AS cliente, v.total,
                           dv.nombre, dv.cantidad, dv.precio, dv.subtotal, dv.ganancia
                    FROM ventas v
                    LEFT JOIN clientes c ON c.id = v.cliente_id
                    LEFT JOIN detalle_venta dv ON v.id = dv.venta_id
                    ORDER BY v.id DESC
                """)
                filas = cursor.fetchall()
            with open(ruta_completa, mode="w", newline="", encoding="utf-8") as archivo_csv:
                escritor = csv.writer(archivo_csv)
                escritor.writerow([
                    "ID_Venta", "Fecha_Venta", "Cliente", "Total_Venta",
                    "Producto", "Cantidad", "Precio_Unitario", "Subtotal_Producto", "Ganancia_Producto"
                ])
                escritor.writerows(filas)
            messagebox.showinfo(
                "✔ Exportación exitosa",
                f"El historial se ha exportado correctamente.\n\nArchivo guardado en:\n{ruta_completa}",
                parent=self
            )
        except Exception as e:
            messagebox.showerror(
                "Error de exportación",
                f"Ocurrió un problema al exportar el archivo CSV:\n{str(e)}",
                parent=self
            )

    # ══════════════════════════════════════════════════════
    #  GESTIÓN DE CONTRASEÑA DE ADMINISTRADOR
    # ══════════════════════════════════════════════════════
    def _verificar_contrasena_inicial(self):
        if get_admin_hash() is not None:
            return
        dlg = tk.Toplevel(self)
        dlg.title("🔐 Crear contraseña de administrador")
        dlg.configure(bg=C["card"])
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.focus_set()
        dlg.protocol("WM_DELETE_WINDOW", lambda: None)
        self.update_idletasks()
        x = self.winfo_x() + (self.winfo_width()  // 2) - 240
        y = self.winfo_y() + (self.winfo_height() // 2) - 160
        dlg.geometry(f"480x320+{x}+{y}")
        tk.Label(dlg, text="🔐  PRIMERA CONFIGURACIÓN",
                 fg=C["accent"], bg=C["card"],
                 font=("Courier", 12, "bold")).pack(pady=(22, 4))
        tk.Label(dlg, text="Crea la contraseña de administrador para proteger\nla eliminación de ventas. Guárdala en un lugar seguro.",
                 fg=C["text"], bg=C["card"],
                 font=("Courier", 9), justify="center").pack(pady=(0, 14))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x", padx=20, pady=(0, 14))
        fields_frame = tk.Frame(dlg, bg=C["card"])
        fields_frame.pack(padx=30, fill="x")
        tk.Label(fields_frame, text="Nueva contraseña:", fg=C["muted"],
                 bg=C["card"], font=("Courier", 9)).pack(anchor="w")
        sv_nueva = tk.StringVar()
        e_nueva = tk.Entry(fields_frame, textvariable=sv_nueva, show="•",
                           bg=C["panel"], fg=C["text"], insertbackground=C["text"],
                           bd=0, font=("Courier", 12), highlightthickness=1,
                           highlightbackground=C["border"])
        e_nueva.pack(fill="x", ipady=7, pady=(2, 10))
        e_nueva.focus()
        tk.Label(fields_frame, text="Confirmar contraseña:", fg=C["muted"],
                 bg=C["card"], font=("Courier", 9)).pack(anchor="w")
        sv_conf = tk.StringVar()
        e_conf = tk.Entry(fields_frame, textvariable=sv_conf, show="•",
                          bg=C["panel"], fg=C["text"], insertbackground=C["text"],
                          bd=0, font=("Courier", 12), highlightthickness=1,
                          highlightbackground=C["border"])
        e_conf.pack(fill="x", ipady=7, pady=(2, 0))
        lbl_error = tk.Label(dlg, text="", fg=C["red"], bg=C["card"],
                             font=("Courier", 9))
        lbl_error.pack(pady=(6, 0))
        def _guardar():
            nueva = sv_nueva.get()
            conf  = sv_conf.get()
            if not nueva:
                lbl_error.config(text="✕ La contraseña no puede estar vacía.")
                return
            if len(nueva) < 4:
                lbl_error.config(text="✕ Mínimo 4 caracteres.")
                return
            if nueva != conf:
                lbl_error.config(text="✕ Las contraseñas no coinciden.")
                sv_conf.set("")
                e_conf.focus()
                return
            set_admin_hash(_hash(nueva))
            dlg.destroy()
            messagebox.showinfo("✔ Contraseña creada",
                "Contraseña de administrador guardada correctamente.\n"
                "Recuérdala: la necesitarás para eliminar ventas.", parent=self)
        e_conf.bind("<Return>", lambda e: _guardar())
        e_nueva.bind("<Return>", lambda e: e_conf.focus())
        tk.Button(dlg, text="✔  Guardar contraseña",
                  bg=C["green"], fg=C["white"], bd=0,
                  font=("Courier", 11, "bold"), pady=10, cursor="hand2",
                  activebackground="#27ae60",
                  command=_guardar).pack(fill="x", padx=30, pady=(12, 0))

    def _cambiar_contrasena(self):
        if get_admin_hash() is None:
            self._verificar_contrasena_inicial()
            return
        dlg = tk.Toplevel(self)
        dlg.title("🔑 Cambiar contraseña de administrador")
        dlg.configure(bg=C["card"])
        dlg.resizable(False, False)
        dlg.grab_set()
        dlg.focus_set()
        dlg.protocol("WM_DELETE_WINDOW", dlg.destroy)
        self.update_idletasks()
        x = self.winfo_x() + (self.winfo_width()  // 2) - 240
        y = self.winfo_y() + (self.winfo_height() // 2) - 180
        dlg.geometry(f"480x360+{x}+{y}")
        tk.Label(dlg, text="🔑  CAMBIAR CONTRASEÑA",
                 fg=C["accent2"], bg=C["card"],
                 font=("Courier", 12, "bold")).pack(pady=(22, 4))
        tk.Label(dlg, text="Debes ingresar tu contraseña actual antes de establecer una nueva.",
                 fg=C["muted"], bg=C["card"],
                 font=("Courier", 9), justify="center").pack(pady=(0, 12))
        tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x", padx=20, pady=(0, 14))
        ff = tk.Frame(dlg, bg=C["card"])
        ff.pack(padx=30, fill="x")
        def _campo(parent, label, sv):
            tk.Label(parent, text=label, fg=C["muted"], bg=C["card"],
                     font=("Courier", 9)).pack(anchor="w")
            ent = tk.Entry(parent, textvariable=sv, show="•",
                           bg=C["panel"], fg=C["text"], insertbackground=C["text"],
                           bd=0, font=("Courier", 12), highlightthickness=1,
                           highlightbackground=C["border"])
            ent.pack(fill="x", ipady=7, pady=(2, 10))
            return ent
        sv_actual = tk.StringVar()
        sv_nueva  = tk.StringVar()
        sv_conf   = tk.StringVar()
        e_actual = _campo(ff, "Contraseña actual:", sv_actual)
        e_nueva  = _campo(ff, "Nueva contraseña:", sv_nueva)
        e_conf   = _campo(ff, "Confirmar nueva contraseña:", sv_conf)
        e_actual.focus()
        lbl_error = tk.Label(dlg, text="", fg=C["red"], bg=C["card"],
                             font=("Courier", 9))
        lbl_error.pack(pady=(0, 4))
        intentos_actuales = [0]
        def _guardar():
            actual = sv_actual.get()
            nueva  = sv_nueva.get()
            conf   = sv_conf.get()
            if not actual or not nueva or not conf:
                lbl_error.config(text="✕ Todos los campos son obligatorios.")
                return
            intentos_actuales[0] += 1
            if _hash(actual) != get_admin_hash():
                restantes = 3 - intentos_actuales[0]
                if restantes <= 0:
                    dlg.destroy()
                    messagebox.showerror("Acceso denegado",
                        "Demasiados intentos incorrectos.\nOperación cancelada.", parent=self)
                    return
                lbl_error.config(
                    text=f"✕ Contraseña actual incorrecta. {restantes} intento(s) restante(s)."
                )
                sv_actual.set("")
                e_actual.focus()
                return
            if len(nueva) < 4:
                lbl_error.config(text="✕ La nueva contraseña debe tener al menos 4 caracteres.")
                return
            if nueva != conf:
                lbl_error.config(text="✕ La nueva contraseña y su confirmación no coinciden.")
                sv_conf.set("")
                e_conf.focus()
                return
            if _hash(nueva) == get_admin_hash():
                lbl_error.config(text="✕ La nueva contraseña es igual a la actual.")
                return
            set_admin_hash(_hash(nueva))
            dlg.destroy()
            messagebox.showinfo("✔ Contraseña actualizada",
                "La contraseña de administrador fue cambiada exitosamente.", parent=self)
        e_actual.bind("<Return>", lambda e: e_nueva.focus())
        e_nueva.bind("<Return>",  lambda e: e_conf.focus())
        e_conf.bind("<Return>",   lambda e: _guardar())
        btn_f = tk.Frame(dlg, bg=C["card"])
        btn_f.pack(padx=30, fill="x", pady=(4, 0))
        tk.Button(btn_f, text="✔  Guardar cambios", bg=C["accent2"], fg=C["white"],
                  bd=0, font=("Courier", 11, "bold"), pady=9, cursor="hand2",
                  command=_guardar).pack(side="left", expand=True, fill="x", padx=(0,4))
        tk.Button(btn_f, text="✕  Cancelar", bg=C["panel"], fg=C["muted"],
                  bd=0, font=("Courier", 10), pady=9, cursor="hand2",
                  command=dlg.destroy).pack(side="left", expand=True, fill="x")

    # ══════════════════════════════════════════════════════
    #  ELIMINAR VENTA CON CONTRASEÑA
    # ══════════════════════════════════════════════════════
    def _eliminar_venta_protegida(self):
        sel = self.tabla_hist.selection()
        if not sel:
            messagebox.showinfo(
                "Sin selección",
                "Primero selecciona una venta de la tabla antes de eliminar.",
                parent=self
            )
            return
        vals = self.tabla_hist.item(sel[0], "values")
        venta_id   = int(vals[0])
        venta_fecha = vals[1]
        venta_total = vals[3]
        password_ok = self._pedir_y_validar_password(venta_id, venta_fecha, venta_total)
        if not password_ok:
            return
        confirmado = messagebox.askyesno(
            "⚠ Confirmar eliminación",
            f"Estás a punto de eliminar PERMANENTEMENTE:\n\n"
            f"  Venta #:  {venta_id}\n"
            f"  Fecha:    {venta_fecha}\n"
            f"  Total:    {venta_total}\n\n"
            "Esta acción NO se puede deshacer.\n¿Continuar?",
            icon="warning",
            parent=self
        )
        if not confirmado:
            return
        try:
            with get_conn() as conn:
                venta_row = conn.execute(
                    "SELECT cliente_id,total FROM ventas WHERE id = ?",
                    (venta_id,)
                ).fetchone()
                conn.execute(
                    "DELETE FROM detalle_venta WHERE venta_id = ?", (venta_id,)
                )
                conn.execute(
                    "DELETE FROM movimientos_caja WHERE venta_id = ?",
                    (venta_id,)
                )
                conn.execute(
                    "DELETE FROM ventas WHERE id = ?", (venta_id,)
                )
                if venta_row and venta_row[0] is not None:
                    cliente_id = int(venta_row[0])
                    total_venta = float(venta_row[1] or 0)
                    conn.execute(
                        "UPDATE clientes"
                        " SET total_compras = CASE"
                        "     WHEN IFNULL(total_compras,0) - ? < 0 THEN 0"
                        "     ELSE IFNULL(total_compras,0) - ?"
                        " END"
                        " WHERE id = ?",
                        (total_venta, total_venta, cliente_id)
                    )
                    ultima = conn.execute(
                        "SELECT MAX(fecha) FROM ventas WHERE cliente_id = ?",
                        (cliente_id,)
                    ).fetchone()[0]
                    conn.execute(
                        "UPDATE clientes SET ultima_visita = ? WHERE id = ?",
                        (ultima, cliente_id)
                    )
        except sqlite3.Error as e:
            messagebox.showerror(
                "Error de base de datos",
                f"No se pudo eliminar la venta.\nDetalle técnico: {e}",
                parent=self
            )
            return
        for row in self.tabla_det.get_children():
            self.tabla_det.delete(row)
        self._cargar_historial()
        if hasattr(self, "tabla_clientes"):
            self._cargar_tabla_clientes()
        self._cargar_clientes_en_venta()
        if hasattr(self, "tabla_conta"):
            self._cargar_tabla_contabilidad()
        messagebox.showinfo(
            "✔ Venta eliminada",
            f"La venta #{venta_id} fue eliminada correctamente.",
            parent=self
        )

    def _pedir_y_validar_password(self, venta_id, fecha, total):
        MAX_INTENTOS = 3
        for intento in range(1, MAX_INTENTOS + 1):
            dlg = tk.Toplevel(self)
            dlg.title("🔒 Autenticación requerida")
            dlg.configure(bg=C["card"])
            dlg.resizable(False, False)
            dlg.grab_set()
            dlg.focus_set()
            self.update_idletasks()
            x = self.winfo_x() + (self.winfo_width()  // 2) - 220
            y = self.winfo_y() + (self.winfo_height() // 2) - 120
            dlg.geometry(f"440x240+{x}+{y}")
            tk.Label(dlg, text="🔒  ELIMINAR VENTA — ACCESO RESTRINGIDO",
                     fg=C["red"], bg=C["card"],
                     font=("Courier", 10, "bold")).pack(pady=(18,4))
            tk.Label(dlg,
                     text=f"Venta #{venta_id}  •  {fecha}  •  {total}",
                     fg=C["muted"], bg=C["card"],
                     font=("Courier", 9)).pack()
            tk.Frame(dlg, bg=C["border"], height=1).pack(fill="x", padx=20, pady=10)
            if intento > 1:
                tk.Label(dlg,
                         text=f"✕ Contraseña incorrecta — intento {intento} de {MAX_INTENTOS}",
                         fg=C["yellow"], bg=C["card"],
                         font=("Courier", 9)).pack(pady=(0,4))
            tk.Label(dlg, text="Ingresa la contraseña de administrador:",
                     fg=C["text"], bg=C["card"],
                     font=("Courier", 10)).pack(pady=(0,6))
            sv_pass = tk.StringVar()
            entry_pass = tk.Entry(dlg, textvariable=sv_pass,
                                  show="•",
                                  bg=C["panel"], fg=C["text"],
                                  insertbackground=C["text"],
                                  bd=0, font=("Courier", 13),
                                  highlightthickness=1,
                                  highlightbackground=C["border"],
                                  width=24)
            entry_pass.pack(ipady=7, pady=(0,12))
            entry_pass.focus()
            resultado = {"accion": None}
            def _confirmar(event=None):
                resultado["accion"] = "ok"
                dlg.destroy()
            def _cancelar(event=None):
                resultado["accion"] = "cancelar"
                dlg.destroy()
            entry_pass.bind("<Return>", _confirmar)
            entry_pass.bind("<Escape>", _cancelar)
            dlg.protocol("WM_DELETE_WINDOW", _cancelar)
            btn_frame = tk.Frame(dlg, bg=C["card"])
            btn_frame.pack()
            tk.Button(btn_frame, text="✔ Confirmar", bg=C["red"], fg=C["white"],
                      bd=0, font=("Courier", 10, "bold"), padx=16, pady=6,
                      cursor="hand2", command=_confirmar).pack(side="left", padx=4)
            tk.Button(btn_frame, text="✕ Cancelar", bg=C["panel"], fg=C["muted"],
                      bd=0, font=("Courier", 10), padx=16, pady=6,
                      cursor="hand2", command=_cancelar).pack(side="left", padx=4)
            dlg.wait_window()
            if resultado["accion"] == "cancelar":
                return False
            ingresada = sv_pass.get()
            if not ingresada:
                messagebox.showwarning(
                    "Campo vacío", "Debes ingresar una contraseña.", parent=self
                )
                continue
            hash_ingresada = _hash(ingresada)
            if hash_ingresada == get_admin_hash():
                return True
            if intento == MAX_INTENTOS:
                messagebox.showerror(
                    "Acceso denegado",
                    f"Se superaron {MAX_INTENTOS} intentos fallidos.\n"
                    "La operación ha sido cancelada por seguridad.",
                    parent=self
                )
                return False
        return False


# ──────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = PuntoDeVenta()
    app.mainloop()
